from __future__ import annotations

import hashlib
import io
from copy import deepcopy
from datetime import datetime, timezone
from typing import Any, Mapping

from backend.profiles.cv_profile_extraction import (
    extract_cv_profile_fallback,
    normalize_profile_payload,
)
from backend.storage.keys import build_private_object_key


CV_EDITOR_SCHEMA_VERSION = "cv_editor_v1"
MAX_EDITOR_HISTORY = 12


class CvEditorRevisionConflict(ValueError):
    def __init__(self, current_revision: int):
        self.current_revision = int(current_revision)
        super().__init__(
            f"This CV changed elsewhere. Reload revision {self.current_revision} before saving again."
        )


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _text(value: Any, *, limit: int = 1200) -> str:
    return " ".join(str(value or "").split())[:limit].strip()


def _multiline(value: Any, *, limit: int = 24) -> list[str]:
    if isinstance(value, list):
        raw_values = value
    else:
        raw_values = str(value or "").replace("\r", "\n").split("\n")
    values: list[str] = []
    seen: set[str] = set()
    for raw_value in raw_values:
        item = " ".join(str(raw_value or "").split()).strip(" -*•")
        if not item or item.casefold() in seen:
            continue
        values.append(item[:600])
        seen.add(item.casefold())
        if len(values) >= limit:
            break
    return values


def _merge_profile_sources(*sources: Mapping[str, Any] | None) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for source in sources:
        if not isinstance(source, Mapping):
            continue
        for key, value in source.items():
            if value not in (None, "", [], {}) and merged.get(str(key)) in (None, "", [], {}):
                merged[str(key)] = deepcopy(value)
    return normalize_profile_payload(merged)


def _editor_profile_from_asset(
    asset: Mapping[str, Any],
    *,
    source_text: str = "",
    shared_profile: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    metadata = dict(asset.get("metadata") or {})
    persisted_editor = metadata.get("cv_editor")
    if isinstance(persisted_editor, Mapping) and isinstance(persisted_editor.get("profile"), Mapping):
        return normalize_profile_payload(persisted_editor["profile"])

    parsed_profile = metadata.get("parsed_profile")
    fallback_profile = extract_cv_profile_fallback(source_text) if source_text.strip() else {}
    return _merge_profile_sources(parsed_profile, fallback_profile, shared_profile)


def build_cv_editor_payload(
    asset: Mapping[str, Any],
    *,
    source_text: str = "",
    shared_profile: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    metadata = dict(asset.get("metadata") or {})
    persisted_editor = metadata.get("cv_editor")
    revision = 0
    if isinstance(persisted_editor, Mapping):
        try:
            revision = max(0, int(persisted_editor.get("revision") or 0))
        except (TypeError, ValueError):
            revision = 0
    return {
        "schema_version": CV_EDITOR_SCHEMA_VERSION,
        "revision": revision,
        "profile": _editor_profile_from_asset(
            asset,
            source_text=source_text,
            shared_profile=shared_profile,
        ),
    }


def profile_to_cv_text(profile: Mapping[str, Any]) -> str:
    normalized = normalize_profile_payload(profile)
    lines: list[str] = []
    identity = [normalized["name"], normalized["role_title"]]
    lines.extend(item for item in identity if item)
    contacts = [
        normalized["email"],
        normalized["location"],
        normalized["website"],
        normalized["linkedin_url"],
        normalized["github_url"],
    ]
    if any(contacts):
        lines.append(" | ".join(item for item in contacts if item))

    def add_section(title: str, values: list[str]) -> None:
        cleaned = [str(value).strip() for value in values if str(value).strip()]
        if not cleaned:
            return
        lines.extend(["", title, *cleaned])

    add_section("Summary", [normalized["summary"]])
    add_section("Skills", normalized["competencies"])
    if normalized["recent_experience"]:
        lines.extend(["", "Experience"])
        for entry in normalized["recent_experience"]:
            heading = " | ".join(
                value
                for value in (_text(entry.get("title")), _text(entry.get("company")), _text(entry.get("period")))
                if value
            )
            if heading:
                lines.append(heading)
            lines.extend(f"- {bullet}" for bullet in entry.get("bullets") or [])
    if normalized["education"]:
        lines.extend(["", "Education"])
        for entry in normalized["education"]:
            heading = " | ".join(
                value
                for value in (
                    _text(entry.get("degree_title")),
                    _text(entry.get("institution")),
                    _text(entry.get("period")),
                )
                if value
            )
            if heading:
                lines.append(heading)
            lines.extend(f"- {detail}" for detail in entry.get("details") or [])
    if normalized["projects"]:
        lines.extend(["", "Projects"])
        for entry in normalized["projects"]:
            heading = " | ".join(
                value for value in (_text(entry.get("title")), _text(entry.get("period"))) if value
            )
            if heading:
                lines.append(heading)
            lines.extend(f"- {bullet}" for bullet in entry.get("bullets") or [])
    add_section("Languages", normalized["languages"])
    for section in normalized["custom_sections"]:
        heading = _text(section.get("heading"), limit=120)
        section_lines = _multiline(section.get("lines"), limit=24)
        if heading and section_lines:
            lines.extend(["", heading, *[f"- {line}" for line in section_lines]])
    return "\n".join(lines).strip() + "\n"


def create_cv_docx_bytes(profile: Mapping[str, Any], *, title: str = "") -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:
        raise RuntimeError("python-docx is required to create editable CV files.") from exc

    normalized = normalize_profile_payload(profile)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(30, 41, 59)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run(normalized["name"] or title or "CV")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(23, 50, 77)
    if normalized["role_title"]:
        role = document.add_paragraph()
        role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_run = role.add_run(normalized["role_title"])
        role_run.font.size = Pt(11)
        role_run.font.color.rgb = RGBColor(13, 148, 136)
    contact_values = [
        normalized["email"],
        normalized["location"],
        normalized["website"],
        normalized["linkedin_url"],
        normalized["github_url"],
    ]
    if any(contact_values):
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run("  ·  ".join(value for value in contact_values if value)).font.size = Pt(8.5)

    def add_heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(13, 148, 136)

    def add_body(text: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(3)

    if normalized["summary"]:
        add_heading("Summary")
        add_body(normalized["summary"])
    if normalized["recent_experience"]:
        add_heading("Experience")
        for entry in normalized["recent_experience"]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            run = paragraph.add_run(_text(entry.get("title")) or "Role")
            run.bold = True
            run.font.color.rgb = RGBColor(23, 50, 77)
            meta = " · ".join(value for value in (_text(entry.get("company")), _text(entry.get("period"))) if value)
            if meta:
                paragraph.add_run(f"  {meta}").italic = True
            for bullet in entry.get("bullets") or []:
                document.add_paragraph(bullet, style="List Bullet")
    if normalized["education"]:
        add_heading("Education")
        for entry in normalized["education"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(_text(entry.get("degree_title")) or "Education")
            run.bold = True
            meta = " · ".join(value for value in (_text(entry.get("institution")), _text(entry.get("period"))) if value)
            if meta:
                paragraph.add_run(f"  {meta}").italic = True
            for detail in entry.get("details") or []:
                document.add_paragraph(detail, style="List Bullet")
    if normalized["projects"]:
        add_heading("Projects")
        for entry in normalized["projects"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(_text(entry.get("title")) or "Project")
            run.bold = True
            if _text(entry.get("period")):
                paragraph.add_run(f"  {_text(entry.get('period'))}").italic = True
            for bullet in entry.get("bullets") or []:
                document.add_paragraph(bullet, style="List Bullet")
    if normalized["competencies"]:
        add_heading("Skills")
        add_body(" · ".join(normalized["competencies"]))
    if normalized["languages"]:
        add_heading("Languages")
        add_body(" · ".join(normalized["languages"]))
    for custom_section in normalized["custom_sections"]:
        section_heading = _text(custom_section.get("heading"), limit=120)
        section_lines = _multiline(custom_section.get("lines"), limit=24)
        if section_heading and section_lines:
            add_heading(section_heading)
            for line in section_lines:
                document.add_paragraph(line, style="List Bullet")

    output = io.BytesIO()
    document.core_properties.title = title or normalized["name"] or "CV"
    document.save(output)
    return output.getvalue()


def save_cv_editor_asset(
    asset: Mapping[str, Any],
    *,
    user_id: str,
    payload: Mapping[str, Any],
    object_storage: Any,
    download_url: str,
    now: str | None = None,
) -> tuple[dict[str, Any], dict[str, Any]]:
    current_metadata = dict(asset.get("metadata") or {})
    current_editor = current_metadata.get("cv_editor")
    current_editor_payload = current_editor if isinstance(current_editor, Mapping) else {}
    current_revision = 0
    if isinstance(current_editor, Mapping):
        try:
            current_revision = max(0, int(current_editor.get("revision") or 0))
        except (TypeError, ValueError):
            current_revision = 0
    try:
        base_revision = int(payload.get("base_revision") or 0)
    except (TypeError, ValueError):
        base_revision = 0
    if base_revision != current_revision:
        raise CvEditorRevisionConflict(current_revision)

    raw_profile = payload.get("profile")
    if not isinstance(raw_profile, Mapping):
        raise ValueError("profile is required.")
    profile = normalize_profile_payload(raw_profile)
    if not profile["name"]:
        raise ValueError("A CV name is required.")
    revision = current_revision + 1
    saved_at = now or _now()
    source_text = profile_to_cv_text(profile)
    docx_bytes = create_cv_docx_bytes(
        profile,
        title=str(asset.get("display_name") or profile["name"] or "CV"),
    )
    content_hash = hashlib.sha256(docx_bytes).hexdigest()
    display_name = str(asset.get("display_name") or "CV").strip() or "CV"
    stem = display_name.rsplit(".", 1)[0] if "." in display_name else display_name
    edited_filename = f"{stem}-edited-v{revision}.docx"
    asset_id = str(asset.get("asset_id") or "").strip()
    if not asset_id:
        raise ValueError("asset_id is required.")
    object_key = build_private_object_key(
        namespace="users",
        owner_id=user_id,
        category="workspace_cv",
        object_id=f"{asset_id}-editor-v{revision}",
        filename=edited_filename,
    )
    object_storage.put(
        object_key,
        docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata={"user_id": user_id, "asset_id": asset_id, "source": "cv_editor"},
    )

    previous_file = dict(asset.get("file") or {})
    previous_key = str(previous_file.get("object_key") or "").strip()
    history = current_metadata.get("cv_editor_versions")
    history = list(history) if isinstance(history, list) else []
    if previous_key:
        history.insert(
            0,
            {
                "revision": current_revision,
                "object_key": previous_key,
                "saved_at": str(current_editor_payload.get("updated_at") or current_metadata.get("processed_at") or ""),
            },
        )

    updated_asset = deepcopy(dict(asset))
    updated_asset["display_name"] = edited_filename
    updated_metadata = dict(current_metadata)
    updated_metadata.update(
        {
            "source_text": source_text,
            "source_char_count": len(source_text),
            "content_sha256": content_hash,
            "status": "ready",
            "parsed_profile": profile,
            "cv_editor": {
                "schema_version": CV_EDITOR_SCHEMA_VERSION,
                "revision": revision,
                "updated_at": saved_at,
                "profile": profile,
            },
            "cv_editor_versions": history[:MAX_EDITOR_HISTORY],
        }
    )
    updated_asset["metadata"] = updated_metadata
    updated_asset["file"] = {
        **previous_file,
        "path": "",
        "object_key": object_key,
        "download_url": download_url,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "extension": "docx",
    }
    return updated_asset, {
        "schema_version": CV_EDITOR_SCHEMA_VERSION,
        "revision": revision,
        "profile": profile,
        "saved_at": saved_at,
        "download_filename": edited_filename,
    }


__all__ = [
    "CV_EDITOR_SCHEMA_VERSION",
    "CvEditorRevisionConflict",
    "build_cv_editor_payload",
    "create_cv_docx_bytes",
    "profile_to_cv_text",
    "save_cv_editor_asset",
]
