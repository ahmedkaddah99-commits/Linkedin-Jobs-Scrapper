import json
import os
import shutil
import subprocess
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Dict, List

from backend.config.job_seeker import cfg_str, load_job_seeker_config, normalize_windows_env_path
from backend.profiles.cv_text import resolve_runtime_cv_docx_path

from .common import sanitize_filename
from .generation import format_header_location, normalize_output_language


DEFAULT_LANGUAGES = [
    "Arabic \u2014 Native",
    "English \u2014 C1",
    "German \u2014 B1/B2",
]
DEFAULT_CV_FONT = "Calibri"

CV_TEMPLATE_PRESETS = {
    "classic": {
        "id": "classic",
        "label": "Classic",
        "description": "Traditional single-column CV with strong dividers.",
        "layout": "classic",
        "heading_case": "upper",
        "base_font_size": 10.5,
        "header_font_size": 17.0,
        "divider_weight": "6",
        "photo_width": 1.5,
        "top_offset_inches": 0.45,
        "page_margin_inches": 0.18,
    },
    "modern": {
        "id": "modern",
        "label": "Modern",
        "description": "Banner-style header with softer section blocks and stronger hierarchy.",
        "layout": "modern",
        "heading_case": "title",
        "base_font_size": 10.8,
        "header_font_size": 18.0,
        "divider_weight": "10",
        "photo_width": 1.45,
        "top_offset_inches": 0.35,
        "page_margin_inches": 0.24,
    },
    "compact": {
        "id": "compact",
        "label": "Compact",
        "description": "Sidebar-style compact layout for concise one-page applications.",
        "layout": "compact",
        "heading_case": "upper",
        "base_font_size": 10.0,
        "header_font_size": 15.5,
        "divider_weight": "4",
        "photo_width": 1.25,
        "top_offset_inches": 0.35,
        "page_margin_inches": 0.18,
    },
    "europass": {
        "id": "europass",
        "label": "EuroPass-style",
        "description": "Structured label-column layout inspired by Europass conventions.",
        "layout": "europass",
        "heading_case": "title",
        "base_font_size": 10.2,
        "header_font_size": 16.0,
        "divider_weight": "8",
        "photo_width": 1.35,
        "top_offset_inches": 0.4,
        "page_margin_inches": 0.22,
    },
    "plain": {
        "id": "plain",
        "label": "Plain",
        "description": "Photo-free single-column resume with a name rule, classic headings, and compact skills.",
        "layout": "plain_resume",
        "heading_case": "title",
        "base_font_size": 10.5,
        "header_font_size": 28.0,
        "divider_weight": "12",
        "photo_width": 0,
        "top_offset_inches": 0,
        "page_margin_inches": 0.7,
        "supports_photo": False,
    },
}

CV_TEMPLATE_ALIASES = {
    "teal_resume": "plain",
}

CV_COLOR_SCHEMES = {
    "classic_navy": {
        "id": "classic_navy",
        "label": "Classic Navy",
        "primary": "1F3A5F",
        "accent": "2EC4B6",
        "surface": "EAF3FF",
    },
    "ocean_teal": {
        "id": "ocean_teal",
        "label": "Ocean Teal",
        "primary": "006B5F",
        "accent": "14B8A6",
        "surface": "E5FFFB",
    },
    "forest": {
        "id": "forest",
        "label": "Forest",
        "primary": "2F5D50",
        "accent": "8AA06F",
        "surface": "F0F7F3",
    },
    "slate": {
        "id": "slate",
        "label": "Slate",
        "primary": "334155",
        "accent": "60A5FA",
        "surface": "EEF2FF",
    },
    "burgundy": {
        "id": "burgundy",
        "label": "Burgundy",
        "primary": "7C2D12",
        "accent": "EA580C",
        "surface": "FFF1EB",
    },
    "charcoal": {
        "id": "charcoal",
        "label": "Charcoal",
        "primary": "111827",
        "accent": "6B7280",
        "surface": "F3F4F6",
    },
}

CV_FONT_OPTIONS = [
    {"id": "Calibri", "label": "Calibri"},
    {"id": "Arial", "label": "Arial"},
    {"id": "Georgia", "label": "Georgia"},
    {"id": "Cambria", "label": "Cambria"},
    {"id": "Aptos", "label": "Aptos"},
]

ALLOWED_PROFILE_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg"}
_WORD_PDF_EXPORT_UNAVAILABLE = False

_CV_SECTION_LABELS = {
    "English": {
        "professional_summary": "Professional Summary",
        "professional_experience": "Professional Experience",
        "projects": "Projects",
        "skills": "Skills",
        "education": "Education",
        "profile": "Profile",
        "languages": "Languages",
        "links": "Links",
    },
    "German": {
        "professional_summary": "Profil",
        "professional_experience": "Berufserfahrung",
        "projects": "Projekte",
        "skills": "Kompetenzen",
        "education": "Ausbildung",
        "profile": "Profil",
        "languages": "Sprachen",
        "links": "Links",
    },
}


def _cv_section_labels(output_language: str) -> dict[str, str]:
    return _CV_SECTION_LABELS.get(normalize_output_language(output_language), _CV_SECTION_LABELS["English"])


def resolve_profile_image_path(raw_path: str):
    candidate = normalize_windows_env_path(raw_path)
    if not candidate:
        return None
    image_path = Path(candidate)
    if image_path.suffix.lower() not in ALLOWED_PROFILE_IMAGE_SUFFIXES:
        return None
    return image_path if image_path.exists() and image_path.is_file() else None


def resolve_optional_image_path(raw_path: str):
    candidate = normalize_windows_env_path(raw_path)
    if not candidate:
        return None
    image_path = Path(candidate)
    return image_path if image_path.exists() and image_path.is_file() else None


def resolve_assets_profile_png(docs_dir: Path):
    user_config_preferred = Path("user_config") / "_profile_from_cv.png"
    if user_config_preferred.exists() and user_config_preferred.is_file():
        return user_config_preferred

    assets_dir = docs_dir / "_assets"
    preferred = assets_dir / "_profile_from_cv.png"
    if preferred.exists() and preferred.is_file():
        return preferred
    user_config_png_files = sorted((Path("user_config")).glob("*.png")) if Path("user_config").exists() else []
    if user_config_png_files:
        return user_config_png_files[0]
    png_files = sorted(assets_dir.glob("*.png")) if assets_dir.exists() else []
    return png_files[0] if png_files else None


def get_document_design_options() -> dict:
    return {
        "templates": list(CV_TEMPLATE_PRESETS.values()),
        "color_schemes": list(CV_COLOR_SCHEMES.values()),
        "fonts": list(CV_FONT_OPTIONS),
    }


def normalize_cv_template_id(template_id: str) -> str:
    normalized_id = str(template_id or "").strip().lower()
    return CV_TEMPLATE_ALIASES.get(normalized_id, normalized_id)


def _resolve_template(template_id: str) -> dict:
    canonical_id = normalize_cv_template_id(template_id)
    return dict(CV_TEMPLATE_PRESETS.get(canonical_id) or CV_TEMPLATE_PRESETS["classic"])


def _resolve_color_scheme(color_scheme_id: str) -> dict:
    raw_value = str(color_scheme_id or "").strip()
    if raw_value in CV_COLOR_SCHEMES:
        return dict(CV_COLOR_SCHEMES[raw_value])
    if len(raw_value.replace("#", "")) == 6:
        normalized = raw_value.replace("#", "").upper()
        return {
            "id": "custom",
            "label": "Custom",
            "primary": normalized,
            "accent": normalized,
            "surface": "F4F7FB",
        }
    return dict(CV_COLOR_SCHEMES["classic_navy"])


def find_cv_docx_source_path():
    runtime_docx_path = resolve_runtime_cv_docx_path()
    if runtime_docx_path is not None:
        return runtime_docx_path

    candidates = []
    config = load_job_seeker_config()
    config_cv_path = normalize_windows_env_path(cfg_str(config, ("candidate", "cv_path"), ""))
    config_cv_docx_path = normalize_windows_env_path(cfg_str(config, ("candidate", "cv_docx_path"), ""))
    if config_cv_path:
        candidates.append(Path(config_cv_path))
    if config_cv_docx_path:
        candidates.append(Path(config_cv_docx_path))
    env_cv_path = normalize_windows_env_path(os.getenv("MY_CV_PATH", ""))
    if env_cv_path:
        candidates.append(Path(env_cv_path))
    candidates.append(Path("Ahmed Kaddah CV.docx"))
    candidates.append(Path(r"C:\Users\ahmed\OneDrive\Personal\CV\Ahmed Kaddah CV.docx"))

    for path in candidates:
        if path.exists() and path.is_file() and path.suffix.lower() == ".docx":
            return path
    return None


def extract_profile_image_from_cv_docx(cv_docx_path: Path, docs_dir: Path):
    try:
        docs_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(cv_docx_path, "r") as archive:
            media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
            if not media_files:
                return None
            media_files.sort()
            first_media = media_files[0]
            extension = Path(first_media).suffix or ".png"
            output_path = docs_dir / f"_profile_from_cv{extension}"
            output_path.write_bytes(archive.read(first_media))
            return output_path if output_path.exists() else None
    except Exception:
        return None


def create_cv_document(
    record: Dict,
    docs_dir: Path,
    run_date: str,
    candidate_name: str,
    candidate_email: str,
    cv_font_name: str,
    cv_template_id: str,
    cv_color_scheme: str,
    languages: List[str],
    profile_image_path,
    include_profile_image: bool,
    profile_links: List[Dict[str, str]],
    output_path: str | Path | None = None,
) -> str:
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:
        raise RuntimeError("python-docx is required to create Word files. Install: pip install python-docx") from exc

    job_id = str(record.get("job_id", "unknown"))
    title = str(record.get("title", "Untitled"))
    company = str(record.get("company", "Unknown Company"))
    header_location = format_header_location(record)
    safe_stem = sanitize_filename(f"{candidate_name}_{title}_{company}_{job_id}_CV", max_length=140)

    if output_path:
        cv_path = Path(output_path)
        cv_path.parent.mkdir(parents=True, exist_ok=True)
    else:
        target_dir = docs_dir / run_date
        target_dir.mkdir(parents=True, exist_ok=True)
        cv_path = target_dir / f"{safe_stem}.docx"

    template = _resolve_template(cv_template_id)
    color_scheme = _resolve_color_scheme(cv_color_scheme)
    primary_rgb = RGBColor.from_string(color_scheme["primary"])
    accent_rgb = RGBColor.from_string(color_scheme["accent"])
    layout = str(template.get("layout") or template["id"]).strip().lower()
    output_language = normalize_output_language(record.get("cv_output_language") or "English")
    section_labels = _cv_section_labels(output_language)

    doc = Document()
    for style_name in ("Normal", "Heading 1", "Heading 2", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = cv_font_name
        style.font.size = Pt(float(template["base_font_size"]))
        try:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), cv_font_name)
        except Exception:
            pass

    section = doc.sections[0]
    if layout == "plain_resume":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.7)
        section.right_margin = Inches(0.8)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
    else:
        margin = Inches(float(template.get("page_margin_inches", 0.1)))
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    language_values = [str(line).strip() for line in (languages or DEFAULT_LANGUAGES) if str(line).strip()]
    language_line = ", ".join(language_values)
    contact_parts = [header_location, candidate_email]
    if language_line:
        contact_parts.append(language_line)
    contact_line = " | ".join([part for part in contact_parts if part])
    summary_text = str(record.get("cv_professional_summary") or "").strip()
    experiences = [item for item in (record.get("cv_professional_experience") or []) if isinstance(item, dict)]
    initiatives = [item for item in (record.get("cv_strategic_initiatives") or []) if isinstance(item, dict)]
    skills = [str(skill).strip() for skill in (record.get("cv_skills") or []) if str(skill).strip()]
    education_items = [item for item in (record.get("cv_education") or []) if isinstance(item, dict)]

    def new_paragraph(target):
        paragraphs = getattr(target, "paragraphs", None)
        if paragraphs is not None and len(paragraphs) == 1:
            paragraph = paragraphs[0]
            if not paragraph.text and not paragraph.runs:
                return paragraph
        return target.add_paragraph()

    def configure_paragraph(paragraph, *, before_pt: float = 0, after_pt: float = 0, line_spacing_pt: float | None = None):
        paragraph.paragraph_format.space_before = Pt(before_pt)
        paragraph.paragraph_format.space_after = Pt(after_pt)
        if line_spacing_pt is not None:
            paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)

    def style_run(run, *, size_pt: float | None = None, bold: bool | None = None, color=None):
        run.font.name = cv_font_name
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if color is not None:
            run.font.color.rgb = color

    def set_paragraph_shading(paragraph, fill_hex: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        p_pr.append(shd)

    def set_paragraph_border(paragraph, *, side: str = "bottom", color_hex: str, size: str, space: str = "0") -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), str(space))
        border.set(qn("w:color"), color_hex)
        p_bdr.append(border)
        p_pr.append(p_bdr)

    def set_cell_shading(cell, fill_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)

    def set_cell_border(cell, *, side: str, color_hex: str, size: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)
        borders.append(border)

    def set_cell_margins(cell, **margins) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for side, value in margins.items():
            node = tc_mar.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def add_hyperlink(paragraph, text: str, url: str, font_size_pt: float = 11.0):
        if not (text and url):
            return
        try:
            relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), relationship_id)

            run = OxmlElement("w:r")
            run_properties = OxmlElement("w:rPr")

            fonts = OxmlElement("w:rFonts")
            for font_attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{font_attribute}"), cv_font_name)
            run_properties.append(fonts)

            color = OxmlElement("w:color")
            color.set(qn("w:val"), color_scheme["accent"])
            run_properties.append(color)

            no_underline = OxmlElement("w:u")
            no_underline.set(qn("w:val"), "none")
            run_properties.append(no_underline)

            font_size = OxmlElement("w:sz")
            font_size.set(qn("w:val"), str(int(font_size_pt * 2)))
            run_properties.append(font_size)

            font_size_cs = OxmlElement("w:szCs")
            font_size_cs.set(qn("w:val"), str(int(font_size_pt * 2)))
            run_properties.append(font_size_cs)

            run.append(run_properties)
            text_node = OxmlElement("w:t")
            text_node.text = text
            run.append(text_node)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
        except Exception:
            fallback_run = paragraph.add_run(text)
            style_run(fallback_run, size_pt=font_size_pt, color=accent_rgb)

    valid_profile_links = []
    for item in profile_links or []:
        if not isinstance(item, dict):
            continue
        link_url = str(item.get("url") or "").strip()
        if not link_url:
            continue
        link_text = str(item.get("text") or "").strip()
        if not link_text:
            continue
        valid_profile_links.append(
            {
                "text": link_text,
                "url": link_url,
                "logo_path": item.get("logo_path"),
            }
        )

    def float_picture_right(run, inline_shape, top_offset_inches: float = 0.45):
        inline = inline_shape._inline
        drawing = run._r.xpath("./w:drawing")[0]

        anchor = OxmlElement("wp:anchor")
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "251658240")
        anchor.set("behindDoc", "0")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")
        anchor.set("distT", "0")
        anchor.set("distB", "0")
        anchor.set("distL", "91440")
        anchor.set("distR", "91440")

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        anchor.append(simple_pos)

        position_h = OxmlElement("wp:positionH")
        position_h.set("relativeFrom", "margin")
        align = OxmlElement("wp:align")
        align.text = "right"
        position_h.append(align)
        anchor.append(position_h)

        position_v = OxmlElement("wp:positionV")
        position_v.set("relativeFrom", "margin")
        pos_offset = OxmlElement("wp:posOffset")
        pos_offset.text = str(int(top_offset_inches * 914400))
        position_v.append(pos_offset)
        anchor.append(position_v)

        extent = inline.xpath("./wp:extent")[0]
        anchor.append(deepcopy(extent))

        effect = inline.xpath("./wp:effectExtent")
        if effect:
            anchor.append(deepcopy(effect[0]))
        else:
            effect_extent = OxmlElement("wp:effectExtent")
            effect_extent.set("l", "0")
            effect_extent.set("t", "0")
            effect_extent.set("r", "0")
            effect_extent.set("b", "0")
            anchor.append(effect_extent)

        wrap_square = OxmlElement("wp:wrapSquare")
        wrap_square.set("wrapText", "bothSides")
        anchor.append(wrap_square)

        doc_pr = inline.xpath("./wp:docPr")[0]
        anchor.append(deepcopy(doc_pr))

        frame_pr = inline.xpath("./wp:cNvGraphicFramePr")
        if frame_pr:
            anchor.append(deepcopy(frame_pr[0]))

        graphic = inline.xpath("./a:graphic")[0]
        anchor.append(deepcopy(graphic))

        drawing.remove(inline)
        drawing.append(anchor)

    def insert_inline_photo(paragraph, *, width_inches: float) -> bool:
        if not (include_profile_image and profile_image_path):
            return False
        try:
            image_run = paragraph.add_run()
            image_run.add_picture(str(profile_image_path), width=Inches(float(width_inches)))
            return True
        except Exception:
            return False

    def insert_floating_photo(paragraph) -> bool:
        if not (include_profile_image and profile_image_path):
            return False
        try:
            image_run = paragraph.add_run()
            inline_shape = image_run.add_picture(
                str(profile_image_path),
                width=Inches(float(template["photo_width"])),
            )
            float_picture_right(
                image_run,
                inline_shape,
                top_offset_inches=float(template["top_offset_inches"]),
            )
            return True
        except Exception:
            return False

    def write_name_block(target, *, include_contact: bool = True, compact_contact: bool = False):
        name_paragraph = new_paragraph(target)
        configure_paragraph(name_paragraph, before_pt=0, after_pt=1)
        name_run = name_paragraph.add_run(candidate_name)
        style_run(name_run, size_pt=float(template["header_font_size"]), bold=True, color=primary_rgb)
        if include_contact and contact_line:
            contact_paragraph = new_paragraph(target)
            configure_paragraph(contact_paragraph, before_pt=0, after_pt=2)
            contact_run = contact_paragraph.add_run(contact_line)
            style_run(
                contact_run,
                size_pt=float(template["base_font_size"]) - (0.2 if compact_contact else 0.0),
                color=accent_rgb,
            )

    def write_text_paragraph(target, text: str, *, after_pt: float = 2, size_pt: float | None = None, bold: bool = False):
        paragraph = new_paragraph(target)
        configure_paragraph(paragraph, before_pt=0, after_pt=after_pt)
        run = paragraph.add_run(text)
        style_run(run, size_pt=size_pt or float(template["base_font_size"]), bold=bold)
        return paragraph

    def write_bullet_paragraph(target, text: str, *, compact: bool = False):
        bullet_text = str(text or "").strip()
        if not bullet_text:
            return None
        paragraph = new_paragraph(target)
        configure_paragraph(paragraph, before_pt=0, after_pt=0 if compact else 1)
        paragraph.paragraph_format.left_indent = Inches(0.2 if compact else 0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        run = paragraph.add_run(f"\u2022 {bullet_text}")
        style_run(run, size_pt=float(template["base_font_size"]) - (0.1 if compact else 0.0))
        return paragraph

    def write_profile_links(target, *, font_size_pt: float = 10.2):
        if not valid_profile_links:
            return
        links_paragraph = new_paragraph(target)
        configure_paragraph(links_paragraph, before_pt=0, after_pt=2)
        icon_size_inches = 16 / 96.0
        for link_index, link_item in enumerate(valid_profile_links):
            if link_index > 0:
                links_paragraph.add_run("   ")
            logo_path = link_item.get("logo_path")
            if logo_path:
                try:
                    logo_run = links_paragraph.add_run()
                    logo_run.add_picture(str(logo_path), width=Inches(icon_size_inches), height=Inches(icon_size_inches))
                    links_paragraph.add_run(" ")
                except Exception:
                    pass
            add_hyperlink(
                links_paragraph,
                text=link_item["text"],
                url=link_item["url"],
                font_size_pt=font_size_pt,
            )

    def add_section_separator(target, *, color_hex: str | None = None, size: str | None = None) -> None:
        sep = new_paragraph(target)
        configure_paragraph(sep, before_pt=0, after_pt=0, line_spacing_pt=1)
        set_paragraph_border(
            sep,
            color_hex=color_hex or color_scheme["primary"],
            size=size or str(template["divider_weight"]),
        )

    def add_standard_section_heading(target, text: str, *, variant: str):
        paragraph = new_paragraph(target)
        configure_paragraph(paragraph, before_pt=1, after_pt=2)
        heading_text = text.upper() if template["heading_case"] == "upper" else text.title()
        run = paragraph.add_run(heading_text)
        style_run(run, size_pt=max(float(template["base_font_size"]) + 1.4, 11.0), bold=True, color=primary_rgb)
        if variant == "modern":
            set_paragraph_shading(paragraph, color_scheme["surface"])
            set_paragraph_border(paragraph, color_hex=color_scheme["accent"], size="4", space="1")
        elif variant == "compact":
            set_paragraph_border(paragraph, color_hex=color_scheme["accent"], size="4", space="0")

    def render_experience(target, *, compact: bool = False):
        for index, item in enumerate(experiences):
            role_title = str(item.get("role_title") or "").strip()
            exp_company = str(item.get("company") or "").strip()
            period = str(item.get("period") or "").strip()
            headline_parts = [part for part in [role_title, exp_company, period] if part]
            if headline_parts:
                exp_header = new_paragraph(target)
                configure_paragraph(exp_header, before_pt=0, after_pt=0 if compact else 1)
                header_run = exp_header.add_run(" | ".join(headline_parts))
                style_run(header_run, bold=True, color=primary_rgb)
            for bullet in item.get("bullets", []):
                bullet_text = str(bullet).strip()
                if not bullet_text:
                    continue
                write_bullet_paragraph(target, bullet_text, compact=compact)
            if index < len(experiences) - 1:
                spacer = new_paragraph(target)
                configure_paragraph(spacer, before_pt=0, after_pt=1 if compact else 2)

    def render_projects(target, *, compact: bool = False):
        for index, item in enumerate(initiatives):
            initiative_title = str(item.get("title") or "").strip()
            if initiative_title:
                ini_header = new_paragraph(target)
                configure_paragraph(ini_header, before_pt=0, after_pt=0 if compact else 1)
                header_run = ini_header.add_run(initiative_title)
                style_run(header_run, bold=True, color=primary_rgb)
            for bullet in item.get("bullets", []):
                bullet_text = str(bullet).strip()
                if not bullet_text:
                    continue
                write_bullet_paragraph(target, bullet_text, compact=compact)
            if index < len(initiatives) - 1:
                spacer = new_paragraph(target)
                configure_paragraph(spacer, before_pt=0, after_pt=1 if compact else 2)

    def render_education(target, *, compact: bool = False):
        for index, item in enumerate(education_items):
            degree_title = str(item.get("degree_title") or "").strip()
            thesis_title = str(item.get("thesis_title") or "").strip()
            if degree_title:
                degree_paragraph = new_paragraph(target)
                configure_paragraph(degree_paragraph, before_pt=0, after_pt=0 if compact else 1)
                degree_run = degree_paragraph.add_run(degree_title)
                style_run(degree_run, bold=True, color=primary_rgb)
            if thesis_title:
                thesis_paragraph = new_paragraph(target)
                configure_paragraph(thesis_paragraph, before_pt=0, after_pt=0 if compact else 1)
                thesis_paragraph.add_run(thesis_title)
            for bullet in item.get("thesis_bullets", []):
                bullet_text = str(bullet).strip()
                if not bullet_text:
                    continue
                write_bullet_paragraph(target, bullet_text, compact=compact)
            if index < len(education_items) - 1:
                spacer = new_paragraph(target)
                configure_paragraph(spacer, before_pt=0, after_pt=1 if compact else 2)

    def render_standard_sections(target, *, variant: str, include_summary: bool = True, include_skills: bool = True):
        rendered_sections = 0

        def begin_section(title: str) -> None:
            nonlocal rendered_sections
            if rendered_sections > 0:
                if variant == "classic":
                    add_section_separator(target)
                else:
                    spacer = new_paragraph(target)
                    configure_paragraph(spacer, before_pt=0, after_pt=2)
            add_standard_section_heading(target, title, variant=variant)
            rendered_sections += 1

        if include_summary and summary_text:
            begin_section(section_labels["professional_summary"])
            write_text_paragraph(target, summary_text, after_pt=2)

        if experiences:
            begin_section(section_labels["professional_experience"])
            render_experience(target, compact=variant == "compact")

        if initiatives:
            begin_section(section_labels["projects"])
            render_projects(target, compact=variant == "compact")

        if include_skills and skills:
            begin_section(section_labels["skills"])
            write_text_paragraph(target, ", ".join(skills), after_pt=2)

        if education_items:
            begin_section(section_labels["education"])
            render_education(target, compact=variant == "compact")

    def add_sidebar_heading(target, text: str):
        paragraph = new_paragraph(target)
        configure_paragraph(paragraph, before_pt=2, after_pt=1)
        run = paragraph.add_run(text.upper())
        style_run(run, size_pt=9.4, bold=True, color=primary_rgb)

    def render_classic_like():
        write_name_block(doc, include_contact=True)
        insert_floating_photo(doc.paragraphs[0])
        write_profile_links(doc, font_size_pt=10.0)
        render_standard_sections(doc, variant="classic")

    def render_modern():
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(5.65)
        header_table.columns[1].width = Inches(1.25)
        left_cell, right_cell = header_table.rows[0].cells
        left_cell.width = Inches(5.65)
        right_cell.width = Inches(1.25)
        set_cell_shading(left_cell, color_scheme["surface"])
        set_cell_shading(right_cell, color_scheme["surface"])

        write_name_block(left_cell, include_contact=True)
        write_profile_links(left_cell, font_size_pt=9.8)
        if include_profile_image and profile_image_path:
            photo_paragraph = right_cell.paragraphs[0]
            configure_paragraph(photo_paragraph, before_pt=0, after_pt=0)
            insert_inline_photo(photo_paragraph, width_inches=float(template["photo_width"]))

        spacer = doc.add_paragraph()
        configure_paragraph(spacer, before_pt=0, after_pt=3)
        render_standard_sections(doc, variant="modern")

    def render_compact():
        write_name_block(doc, include_contact=True, compact_contact=True)
        body_table = doc.add_table(rows=1, cols=2)
        body_table.autofit = False
        body_table.columns[0].width = Inches(2.0)
        body_table.columns[1].width = Inches(4.9)
        sidebar_cell, main_cell = body_table.rows[0].cells
        sidebar_cell.width = Inches(2.0)
        main_cell.width = Inches(4.9)
        set_cell_shading(sidebar_cell, color_scheme["surface"])

        if include_profile_image and profile_image_path:
            photo_paragraph = sidebar_cell.paragraphs[0]
            configure_paragraph(photo_paragraph, before_pt=0, after_pt=3)
            insert_inline_photo(photo_paragraph, width_inches=max(float(template["photo_width"]), 1.45))

        if summary_text:
            add_sidebar_heading(sidebar_cell, section_labels["profile"])
            write_text_paragraph(sidebar_cell, summary_text, after_pt=2, size_pt=float(template["base_font_size"]) - 0.2)

        if skills:
            add_sidebar_heading(sidebar_cell, section_labels["skills"])
            for skill in skills:
                write_text_paragraph(sidebar_cell, skill, after_pt=0.5, size_pt=float(template["base_font_size"]) - 0.1)

        if language_values:
            add_sidebar_heading(sidebar_cell, section_labels["languages"])
            for language in language_values:
                write_text_paragraph(sidebar_cell, language, after_pt=0.5, size_pt=float(template["base_font_size"]) - 0.1)

        if valid_profile_links:
            add_sidebar_heading(sidebar_cell, section_labels["links"])
            write_profile_links(sidebar_cell, font_size_pt=9.4)

        render_standard_sections(
            main_cell,
            variant="compact",
            include_summary=False,
            include_skills=False,
        )

    def render_europass():
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(5.55)
        header_table.columns[1].width = Inches(1.35)
        left_cell, right_cell = header_table.rows[0].cells
        left_cell.width = Inches(5.55)
        right_cell.width = Inches(1.35)
        write_name_block(left_cell, include_contact=True)
        write_profile_links(left_cell, font_size_pt=9.8)
        if include_profile_image and profile_image_path:
            photo_paragraph = right_cell.paragraphs[0]
            configure_paragraph(photo_paragraph, before_pt=0, after_pt=0)
            insert_inline_photo(photo_paragraph, width_inches=float(template["photo_width"]))

        def render_europass_section(title: str, content_writer):
            section_table = doc.add_table(rows=1, cols=2)
            section_table.autofit = False
            section_table.columns[0].width = Inches(1.55)
            section_table.columns[1].width = Inches(5.35)
            label_cell, body_cell = section_table.rows[0].cells
            label_cell.width = Inches(1.55)
            body_cell.width = Inches(5.35)
            set_cell_shading(label_cell, color_scheme["surface"])
            label_paragraph = label_cell.paragraphs[0]
            configure_paragraph(label_paragraph, before_pt=0, after_pt=0)
            label_run = label_paragraph.add_run(title.title())
            style_run(label_run, size_pt=9.6, bold=True, color=primary_rgb)
            content_writer(body_cell)

        if summary_text:
            render_europass_section(section_labels["profile"], lambda target: write_text_paragraph(target, summary_text, after_pt=1))
        if experiences:
            render_europass_section(section_labels["professional_experience"], lambda target: render_experience(target, compact=True))
        if initiatives:
            render_europass_section(section_labels["projects"], lambda target: render_projects(target, compact=True))
        if skills:
            render_europass_section(section_labels["skills"], lambda target: write_text_paragraph(target, ", ".join(skills), after_pt=1))
        if education_items:
            render_europass_section(section_labels["education"], lambda target: render_education(target, compact=True))
        if language_values:
            render_europass_section(section_labels["languages"], lambda target: write_text_paragraph(target, language_line, after_pt=1))

    def render_plain_resume():
        header_table = doc.add_table(rows=1, cols=1)
        header_cell = header_table.cell(0, 0)
        set_cell_border(
            header_cell,
            side="bottom",
            color_hex=color_scheme["accent"],
            size=str(template["divider_weight"]),
        )
        set_cell_margins(header_cell, top=0, right=0, bottom=40, left=0)
        name_paragraph = header_cell.paragraphs[0]
        configure_paragraph(name_paragraph, before_pt=0, after_pt=0)
        name_run = name_paragraph.add_run(candidate_name)
        style_run(
            name_run,
            size_pt=float(template["header_font_size"]),
            color=primary_rgb,
        )

        contact_paragraph = doc.add_paragraph()
        configure_paragraph(contact_paragraph, before_pt=6, after_pt=0)
        has_contact_item = False
        if contact_line:
            contact_run = contact_paragraph.add_run(contact_line)
            style_run(contact_run, size_pt=10.0)
            has_contact_item = True
        for link in valid_profile_links:
            if has_contact_item:
                separator = contact_paragraph.add_run(" | ")
                style_run(separator, size_pt=10.0)
            add_hyperlink(
                contact_paragraph,
                text=link["text"],
                url=link["url"],
                font_size_pt=10.0,
            )
            has_contact_item = True

        def begin_section(title: str) -> None:
            heading = doc.add_paragraph()
            configure_paragraph(heading, before_pt=12, after_pt=4)
            run = heading.add_run(title)
            style_run(run, size_pt=14.0, bold=True, color=primary_rgb)

        def render_plain_experience() -> None:
            for item in experiences:
                headline_parts = [
                    str(item.get("role_title") or "").strip(),
                    str(item.get("company") or "").strip(),
                    str(item.get("period") or "").strip(),
                ]
                headline = " | ".join(part for part in headline_parts if part)
                if headline:
                    paragraph = doc.add_paragraph()
                    configure_paragraph(paragraph, before_pt=3, after_pt=2)
                    run = paragraph.add_run(headline.upper())
                    style_run(run, size_pt=12.0, bold=True)
                for bullet in item.get("bullets", []):
                    write_bullet_paragraph(doc, str(bullet).strip(), compact=True)

        def render_plain_projects() -> None:
            for item in initiatives:
                title_text = str(item.get("title") or "").strip()
                if title_text:
                    paragraph = doc.add_paragraph()
                    configure_paragraph(paragraph, before_pt=3, after_pt=2)
                    run = paragraph.add_run(title_text.upper())
                    style_run(run, size_pt=12.0, bold=True)
                for bullet in item.get("bullets", []):
                    write_bullet_paragraph(doc, str(bullet).strip(), compact=True)

        def render_plain_education() -> None:
            for item in education_items:
                degree_title = str(item.get("degree_title") or "").strip()
                thesis_title = str(item.get("thesis_title") or "").strip()
                if degree_title:
                    paragraph = doc.add_paragraph()
                    configure_paragraph(paragraph, before_pt=3, after_pt=2)
                    run = paragraph.add_run(degree_title.upper())
                    style_run(run, size_pt=12.0, bold=True)
                if thesis_title:
                    write_text_paragraph(doc, thesis_title, after_pt=1)
                for bullet in item.get("thesis_bullets", []):
                    write_bullet_paragraph(doc, str(bullet).strip(), compact=True)

        def render_plain_skills() -> None:
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph, before_pt=0, after_pt=1, line_spacing_pt=12.5)
            run = paragraph.add_run(" | ".join(skills))
            style_run(run, size_pt=float(template["base_font_size"]))

        if summary_text:
            begin_section(section_labels["profile"])
            write_text_paragraph(doc, summary_text, after_pt=2)
        if experiences:
            begin_section(section_labels["professional_experience"].replace("Professional ", ""))
            render_plain_experience()
        if initiatives:
            begin_section(section_labels["projects"])
            render_plain_projects()
        if education_items:
            begin_section(section_labels["education"])
            render_plain_education()
        if skills:
            begin_section("Skills & Abilities" if output_language == "English" else section_labels["skills"])
            render_plain_skills()

    if layout == "modern":
        render_modern()
    elif layout == "compact":
        render_compact()
    elif layout == "europass":
        render_europass()
    elif layout == "plain_resume":
        render_plain_resume()
    else:
        render_classic_like()

    doc.save(cv_path)
    return str(cv_path.resolve())


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _profile_link_urls(profile_links: List[Dict[str, str]]) -> dict:
    urls = {"website": "", "linkedin_url": "", "github_url": ""}
    for item in profile_links or []:
        if not isinstance(item, dict):
            continue
        url = str(item.get("url") or "").strip()
        if not url:
            continue
        text = str(item.get("text") or item.get("icon") or "").strip().lower()
        normalized_url = url.lower()
        if "linkedin" in text or "linkedin.com" in normalized_url:
            urls["linkedin_url"] = urls["linkedin_url"] or url
        elif "github" in text or "github.com" in normalized_url:
            urls["github_url"] = urls["github_url"] or url
        else:
            urls["website"] = urls["website"] or url
    return urls


def build_cv_html_export_payload(
    record: Dict,
    *,
    candidate_name: str,
    candidate_email: str,
    cv_font_name: str,
    cv_template_id: str,
    cv_color_scheme: str,
    languages: List[str],
    profile_image_path,
    include_profile_image: bool,
    profile_links: List[Dict[str, str]],
) -> dict:
    link_urls = _profile_link_urls(profile_links)
    cv_output_language = normalize_output_language(record.get("cv_output_language") or "English")
    profile = {
        "name": str(candidate_name or "").strip(),
        "role_title": str(record.get("title") or "").strip(),
        "target_role": str(record.get("title") or "").strip(),
        "target_company": str(record.get("company") or "").strip(),
        "location": format_header_location(record),
        "email": str(candidate_email or "").strip(),
        "website": link_urls["website"],
        "linkedin_url": link_urls["linkedin_url"],
        "github_url": link_urls["github_url"],
        "summary": str(record.get("cv_professional_summary") or "").strip(),
        "recent_experience": [
            {
                "title": str(item.get("role_title") or item.get("title") or "").strip(),
                "company": str(item.get("company") or "").strip(),
                "period": str(item.get("period") or "").strip(),
                "bullets": [str(bullet).strip() for bullet in item.get("bullets", []) if str(bullet).strip()],
            }
            for item in (record.get("cv_professional_experience") or [])
            if isinstance(item, dict)
        ],
        "projects": [
            {
                "title": str(item.get("title") or item.get("name") or "").strip(),
                "period": str(item.get("period") or item.get("date") or item.get("year") or "").strip(),
                "bullets": [str(bullet).strip() for bullet in item.get("bullets", []) if str(bullet).strip()],
            }
            for item in (record.get("cv_strategic_initiatives") or [])
            if isinstance(item, dict)
        ],
        "competencies": [str(skill).strip() for skill in (record.get("cv_skills") or []) if str(skill).strip()],
        "languages": [str(line).strip() for line in (languages or []) if str(line).strip()],
        "education": [item for item in (record.get("cv_education") or []) if isinstance(item, dict)],
        "cv_output_language": cv_output_language,
    }
    documents = {
        "cv_template": str(cv_template_id or "classic").strip() or "classic",
        "cv_color_scheme": str(cv_color_scheme or "classic_navy").strip() or "classic_navy",
        "cv_font": str(cv_font_name or DEFAULT_CV_FONT).strip() or DEFAULT_CV_FONT,
        "include_photo": bool(include_profile_image),
        "cv_output_language": cv_output_language,
    }
    payload = {
        "profile": profile,
        "documents": documents,
        "workspaceSettings": {"cv_output_language": cv_output_language},
    }
    if include_profile_image and profile_image_path:
        payload["photo_path"] = str(profile_image_path)
    return payload


def create_cv_pdf_document(
    record: Dict,
    *,
    output_path: str | Path,
    candidate_name: str,
    candidate_email: str,
    cv_font_name: str,
    cv_template_id: str,
    cv_color_scheme: str,
    languages: List[str],
    profile_image_path,
    include_profile_image: bool,
    profile_links: List[Dict[str, str]],
) -> str:
    renderer_script = _repo_root() / "frontend" / "scripts" / "render-cv-pdf.mjs"
    if not renderer_script.exists():
        raise RuntimeError(f"CV HTML PDF renderer not found: {renderer_script}")

    target_path = Path(output_path)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    payload = build_cv_html_export_payload(
        record,
        candidate_name=candidate_name,
        candidate_email=candidate_email,
        cv_font_name=cv_font_name,
        cv_template_id=cv_template_id,
        cv_color_scheme=cv_color_scheme,
        languages=languages,
        profile_image_path=profile_image_path,
        include_profile_image=include_profile_image,
        profile_links=profile_links,
    )
    command = ["node", str(renderer_script), "--output", str(target_path)]
    try:
        completed = subprocess.run(
            command,
            input=json.dumps(payload),
            check=True,
            capture_output=True,
            text=True,
            timeout=45,
            cwd=str(_repo_root() / "frontend"),
        )
    except FileNotFoundError as exc:
        raise RuntimeError("Unable to render CV PDF because Node.js is not available on PATH.") from exc
    except subprocess.CalledProcessError as exc:
        details = (exc.stderr or exc.stdout or "").strip()
        raise RuntimeError(f"HTML CV PDF export failed. {details}") from exc
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("HTML CV PDF export timed out after 45 seconds.") from exc

    if not target_path.exists():
        details = (completed.stderr or completed.stdout or "").strip()
        raise RuntimeError(f"HTML CV PDF export did not create {target_path}. {details}")
    return str(target_path.resolve())


def convert_docx_to_pdf(docx_path: str) -> str:
    global _WORD_PDF_EXPORT_UNAVAILABLE
    source_path = Path(docx_path)
    target_path = source_path.with_suffix(".pdf")

    office_cmd = shutil.which("soffice") or shutil.which("libreoffice")
    if office_cmd:
        try:
            subprocess.run(
                [office_cmd, "--headless", "--convert-to", "pdf", "--outdir", str(source_path.parent), str(source_path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=20,
            )
            if target_path.exists():
                return str(target_path.resolve())
        except Exception:
            pass

    if not _WORD_PDF_EXPORT_UNAVAILABLE:
        try:
            src = str(source_path.resolve()).replace("'", "''")
            dst = str(target_path.resolve()).replace("'", "''")
            ps_script = (
                "$ErrorActionPreference = 'Stop'; "
                "$word = $null; "
                "$doc = $null; "
                "try { "
                "$word = New-Object -ComObject Word.Application; "
                "$word.Visible = $false; "
                "$word.DisplayAlerts = 0; "
                f"$doc = $word.Documents.Open('{src}', $false, $true); "
                f"$doc.ExportAsFixedFormat('{dst}', 17); "
                "} finally { "
                "if ($doc -ne $null) { $doc.Close($false) }; "
                "if ($word -ne $null) { $word.Quit() } "
                "}"
            )
            subprocess.run(
                ["powershell", "-NoProfile", "-Command", ps_script],
                check=True,
                capture_output=True,
                text=True,
                timeout=20,
            )
            if target_path.exists():
                return str(target_path.resolve())
        except Exception:
            _WORD_PDF_EXPORT_UNAVAILABLE = True

    raise RuntimeError(
        "Unable to convert DOCX to PDF. Install LibreOffice and add soffice to PATH, "
        "or ensure Microsoft Word automation can export PDFs without prompts."
    )
