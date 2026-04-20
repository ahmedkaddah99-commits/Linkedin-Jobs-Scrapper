from __future__ import annotations

import re
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Dict, List, Mapping

from backend.capabilities.tailored_documents.rendering import (
    create_cv_document,
    convert_docx_to_pdf,
    resolve_assets_profile_png,
    resolve_optional_image_path,
    resolve_profile_image_path,
)
from backend.config.job_seeker import (
    cfg_bool as job_cfg_bool,
    cfg_str as job_cfg_str,
    load_job_seeker_config,
    normalize_windows_env_path,
)
from .support import (
    cfg_list,
    cfg_str,
    load_baseline_profile_text,
    load_reusable_packages_config,
    load_json_file,
    resolve_path,
    save_json_file,
)


ROLE_SUMMARY_BY_CATEGORY = {
    "moving_helper_loader": "Reliable and physically resilient operations helper with practical logistics background. Experienced in loading, unloading, and coordinating movement of equipment while maintaining safety and punctual execution.",
    "waiter_service_staff": "Customer-focused service worker with hands-on operations and field-service experience. Strong under pressure, fast in busy environments, and committed to high service quality.",
    "cook_kitchen_staff": "Disciplined kitchen support profile with strong routine execution, hygiene awareness, and ability to keep fast-paced preparation and cleanup tasks on schedule.",
    "warehouse_logistics": "Operations-focused logistics worker with practical fleet and inventory coordination experience. Strong in sorting, handling, and day-to-day warehouse workflows with high reliability.",
    "delivery_driver": "Field operations profile with route coordination and transport workflow experience. Strong time management, reliability, and careful handling of goods in daily delivery operations.",
    "cleaning_facility_support": "Dependable support worker with strong execution discipline, attention to detail, and consistency in maintaining clean, safe, and operational work environments.",
    "production_packaging": "Hands-on production helper profile with strong stamina and focus on throughput, quality checks, and packaging tasks in repetitive and fast-paced environments.",
    "retail_store_support": "Store operations support profile with practical customer-facing and inventory support experience. Reliable in stock handling, floor support, and daily operational routines.",
    "other_operational": "Reliable operations worker with practical logistics and service background, high work ethic, and readiness for physical and shift-based operational tasks.",
}

ROLE_SKILLS_BY_CATEGORY = {
    "moving_helper_loader": ["Loading and unloading support", "Packing and material handling", "Team lifting and safety awareness", "Fast execution in physical tasks"],
    "waiter_service_staff": ["Customer service in high-traffic settings", "Order and table support", "Team coordination during peak hours", "Reliable shift execution"],
    "cook_kitchen_staff": ["Kitchen prep and station support", "Dishwashing and hygiene routines", "Food handling discipline", "Fast support in busy shifts"],
    "warehouse_logistics": ["Warehouse and inventory routines", "Sorting, picking, and packing", "Logistics workflow coordination", "Operational reliability"],
    "delivery_driver": ["Route and dispatch coordination", "Timely delivery workflow support", "Careful goods handling", "Daily service reliability"],
    "cleaning_facility_support": ["Cleaning process consistency", "Facility hygiene support", "Attention to detail", "Shift reliability"],
    "production_packaging": ["Production line support", "Packaging and labeling routines", "Repetitive-task endurance", "Quality and throughput awareness"],
    "retail_store_support": ["Stockroom and shelf support", "Basic cashier/store operations support", "Customer-facing communication", "Daily opening/closing task support"],
    "other_operational": ["Physical work readiness", "Operational discipline", "Team collaboration", "Punctual and reliable execution"],
}


def sanitize_filename(value: str, max_length: int = 80) -> str:
    cleaned = "".join(char if char.isalnum() or char in ("-", "_") else "_" for char in (value or ""))
    while "__" in cleaned:
        cleaned = cleaned.replace("__", "_")
    cleaned = cleaned.strip("_")
    return (cleaned or "item")[:max_length]


def _normalize_marker(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).upper()


def _extract_section_lines(text: str, start_marker: str, end_markers: List[str]) -> List[str]:
    lines = [line.rstrip() for line in (text or "").splitlines()]
    start_index = None
    normalized_start = _normalize_marker(start_marker)
    normalized_end_markers = {_normalize_marker(item) for item in end_markers}

    for index, line in enumerate(lines):
        if _normalize_marker(line) == normalized_start:
            start_index = index + 1
            break

    if start_index is None:
        return []

    collected: List[str] = []
    for line in lines[start_index:]:
        if _normalize_marker(line) in normalized_end_markers:
            break
        collected.append(line.rstrip())
    return collected


def _extract_bullets(lines: List[str]) -> List[str]:
    bullets: List[str] = []
    for raw_line in lines:
        line = str(raw_line or "").strip()
        if not line:
            continue
        if re.match(r"^[\-\*\u2022]\s*", line):
            bullets.append(re.sub(r"^[\-\*\u2022]+\s*", "", line).strip())
        elif bullets:
            bullets[-1] = f"{bullets[-1]} {line}".strip()
    return bullets


def _extract_strength_labels(bullets: List[str]) -> List[str]:
    labels: List[str] = []
    for bullet in bullets:
        cleaned = str(bullet or "").strip()
        if not cleaned:
            continue
        if ":" in cleaned:
            cleaned = cleaned.split(":", 1)[0].strip()
        if cleaned and cleaned not in labels:
            labels.append(cleaned)
    return labels


def _parse_baseline_experiences(text: str) -> List[Dict[str, Any]]:
    lines = _extract_section_lines(text, "BERUFLICHER WERDEGANG", [])
    experiences: List[Dict[str, Any]] = []
    current: Dict[str, Any] | None = None

    for raw_line in lines:
        line = str(raw_line or "").strip()
        if not line:
            continue

        if "|" in line and not re.match(r"^[\-\*\u2022]\s*", line):
            if current is not None:
                experiences.append(current)
            parts = [part.strip() for part in line.split("|")]
            current = {
                "role_title": parts[0] if len(parts) > 0 else "",
                "company": parts[1] if len(parts) > 1 else "",
                "period": parts[2] if len(parts) > 2 else "",
                "bullets": [],
            }
            continue

        bullet = re.sub(r"^[\-\*\u2022]+\s*", "", line).strip()
        if current is None:
            continue
        if re.match(r"^[\-\*\u2022]\s*", line):
            current["bullets"].append(bullet)
        elif current["bullets"]:
            current["bullets"][-1] = f"{current['bullets'][-1]} {line}".strip()

    if current is not None:
        experiences.append(current)

    return experiences


def _parse_baseline_profile(text: str) -> Dict[str, Any]:
    summary_lines = _extract_section_lines(text, "ZUSAMMENFASSUNG", ["STARKE", "BERUFLICHER WERDEGANG"])
    strength_lines = _extract_section_lines(text, "STARKE", ["BERUFLICHER WERDEGANG"])
    strengths = _extract_bullets(strength_lines)
    return {
        "summary": " ".join(line.strip() for line in summary_lines if line.strip()).strip(),
        "strengths": strengths,
        "strength_labels": _extract_strength_labels(strengths),
        "experiences": _parse_baseline_experiences(text),
    }


def _merge_unique_items(*groups: List[str]) -> List[str]:
    merged: List[str] = []
    seen = set()
    for group in groups:
        for item in group or []:
            cleaned = str(item or "").strip()
            if not cleaned:
                continue
            key = cleaned.casefold()
            if key in seen:
                continue
            seen.add(key)
            merged.append(cleaned)
    return merged


def build_role_cv_record(
    candidate_location: str,
    baseline_profile: Dict[str, Any],
    category: Dict,
) -> Dict[str, Any]:
    category_id = str(category.get("id") or "other_operational")
    category_name = str(category.get("name") or "Other Operational")
    role_skills = ROLE_SKILLS_BY_CATEGORY.get(category_id, ROLE_SKILLS_BY_CATEGORY["other_operational"])
    baseline_summary = str(baseline_profile.get("summary") or "").strip()
    summary_parts = [baseline_summary] if baseline_summary else []
    summary_parts.append(f"Zielprofil mit Fokus auf den Bereich {category_name}.")
    skills = _merge_unique_items(baseline_profile.get("strength_labels", []), role_skills)
    experiences = [dict(item) for item in baseline_profile.get("experiences", []) if isinstance(item, dict)]
    return {
        "job_id": category_id,
        "title": category_name,
        "company": "Reusable Role Profile",
        "location_raw": candidate_location,
        "cv_professional_summary": " ".join(part for part in summary_parts if part).strip(),
        "cv_professional_experience": experiences,
        "cv_strategic_initiatives": [],
        "cv_skills": skills,
        "cv_education": [],
    }


def build_role_cv_text(
    candidate_name: str,
    candidate_email: str,
    candidate_location: str,
    availability: str,
    languages: List[str],
    record: Dict[str, Any],
) -> str:
    lines = [
        candidate_name,
        f"Email: {candidate_email}",
        f"Location: {candidate_location}",
        f"Availability: {availability}",
    ]
    if languages:
        lines.append(f"Languages: {', '.join(languages)}")

    title = str(record.get("title") or "").strip()
    if title:
        lines.extend(["", f"Target Role: {title}"])

    summary = str(record.get("cv_professional_summary") or "").strip()
    if summary:
        lines.extend(["", "Professional Summary", summary])

    experiences = record.get("cv_professional_experience") or []
    if experiences:
        lines.extend(["", "Professional Experience"])
        for item in experiences:
            if not isinstance(item, dict):
                continue
            header_parts = [
                str(item.get("role_title") or "").strip(),
                str(item.get("company") or "").strip(),
                str(item.get("period") or "").strip(),
            ]
            header = " | ".join(part for part in header_parts if part)
            if header:
                lines.append(header)
            for bullet in item.get("bullets", []):
                cleaned = str(bullet or "").strip()
                if cleaned:
                    lines.append(f"- {cleaned}")

    skills = [str(item).strip() for item in (record.get("cv_skills") or []) if str(item).strip()]
    if skills:
        lines.extend(["", "Skills"])
        lines.extend([f"- {item}" for item in skills])

    return "\n".join(lines).strip()


def write_docx(path: Path, text: str) -> bool:
    try:
        from docx import Document
    except Exception:
        return False
    document = Document()
    for raw_line in (text or "").splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        document.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return True


def build_stage4_args(config: dict | None = None, overrides: Mapping[str, Any] | None = None) -> SimpleNamespace:
    config = config or load_reusable_packages_config()
    style_config = load_job_seeker_config()
    profile_links = []
    for key, default_text, default_icon in (("linkedin", "LinkedIn", "in"), ("github", "GitHub", "GH")):
        profile_links.append(
            {
                "icon": job_cfg_str(style_config, ("candidate", "profile_links", key, "icon"), default_icon),
                "text": job_cfg_str(style_config, ("candidate", "profile_links", key, "text"), default_text),
                "url": job_cfg_str(style_config, ("candidate", "profile_links", key, "url"), ""),
                "logo_path": job_cfg_str(style_config, ("candidate", "profile_links", key, "logo_path"), ""),
            }
        )
    payload = {
        "input": cfg_str(config, ("runtime", "stage4", "input_json"), "outputs/stage3_classified_jobs.json"),
        "role_cv_output_dir": cfg_str(config, ("runtime", "stage4", "role_cv_output_dir"), "outputs/role_cvs"),
        "role_cv_index_json": cfg_str(config, ("runtime", "stage4", "role_cv_index_json"), "outputs/stage4_role_cvs.json"),
        "categories": cfg_list(config, ("classification", "categories"), []),
        "candidate_name": cfg_str(config, ("candidate", "name"), "Ahmed Kaddah"),
        "candidate_email": cfg_str(config, ("candidate", "email"), ""),
        "candidate_location": cfg_str(config, ("candidate", "location"), ""),
        "availability": cfg_str(config, ("candidate", "availability"), ""),
        "languages": [str(item) for item in cfg_list(config, ("candidate", "languages"), []) if str(item).strip()],
        "profile_image": job_cfg_str(style_config, ("candidate", "profile_image_path"), ""),
        "include_photo": job_cfg_bool(style_config, ("candidate", "include_photo"), True),
        "cv_font": job_cfg_str(style_config, ("candidate", "cv_font"), "Calibri"),
        "cv_template": job_cfg_str(style_config, ("candidate", "cv_template"), "classic") or "classic",
        "cv_color_scheme": job_cfg_str(style_config, ("candidate", "cv_color_scheme"), "classic_navy")
        or "classic_navy",
        "profile_links": profile_links,
        "generate_pdf": True,
    }
    if overrides:
        payload.update({key: value for key, value in overrides.items() if value is not None})
    return SimpleNamespace(**payload)


def run_stage4_pipeline(args, *, config: dict | None = None, jobs: List[Dict] | None = None) -> dict[str, Any]:
    _ = config
    input_path = resolve_path(args.input)
    if jobs is None:
        if not input_path.exists():
            raise FileNotFoundError(f"input file not found: {input_path}")
        jobs = load_json_file(input_path)
        if not isinstance(jobs, list):
            raise ValueError("input JSON must be a list of jobs.")

    categories = [item for item in (getattr(args, "categories", None) or []) if isinstance(item, dict) and item.get("id")]
    if not categories:
        raise ValueError("no categories configured.")

    jobs_per_category: Dict[str, int] = {}
    for job in jobs:
        category_id = str(job.get("role_category_id") or "other_operational")
        jobs_per_category[category_id] = jobs_per_category.get(category_id, 0) + 1

    role_cv_output_dir = resolve_path(args.role_cv_output_dir)
    baseline_cv_text = load_baseline_profile_text()
    baseline_profile = _parse_baseline_profile(baseline_cv_text)
    cv_index_records: List[Dict] = []
    normalized_profile_input = normalize_windows_env_path(str(getattr(args, "profile_image", "") or ""))
    include_photo = bool(getattr(args, "include_photo", True))
    profile_image_path = resolve_profile_image_path(normalized_profile_input) if include_photo else None
    if include_photo and not profile_image_path:
        profile_image_path = resolve_assets_profile_png(role_cv_output_dir.parent)

    profile_links = []
    for item in getattr(args, "profile_links", []) or []:
        if not isinstance(item, dict):
            continue
        url = str(item.get("url") or "").strip()
        text_value = str(item.get("text") or "").strip()
        if not url or not text_value:
            continue
        profile_links.append(
            {
                "icon": str(item.get("icon") or "").strip(),
                "text": text_value,
                "url": url,
                "logo_path": resolve_optional_image_path(str(item.get("logo_path") or "").strip()),
            }
        )

    for category in categories:
        category_id = str(category.get("id") or "").strip()
        category_name = str(category.get("name") or "").strip() or category_id
        if not category_id:
            continue
        record = build_role_cv_record(
            candidate_location=args.candidate_location,
            baseline_profile=baseline_profile,
            category=category,
        )
        cv_text = build_role_cv_text(
            candidate_name=args.candidate_name,
            candidate_email=args.candidate_email,
            candidate_location=args.candidate_location,
            availability=args.availability,
            languages=list(args.languages),
            record=record,
        )
        safe_name = sanitize_filename(f"{category_id}_{category_name}")
        txt_path = role_cv_output_dir / f"{safe_name}.txt"
        docx_path = role_cv_output_dir / f"{safe_name}.docx"
        pdf_path = role_cv_output_dir / f"{safe_name}.pdf"
        txt_path.parent.mkdir(parents=True, exist_ok=True)
        txt_path.write_text(cv_text, encoding="utf-8")
        docx_created = False
        pdf_created = False
        docx_error = ""
        pdf_generation_error = ""
        try:
            create_cv_document(
                record,
                docs_dir=role_cv_output_dir,
                run_date="",
                candidate_name=args.candidate_name,
                candidate_email=args.candidate_email,
                cv_font_name=str(getattr(args, "cv_font", "Calibri") or "Calibri"),
                cv_template_id=str(getattr(args, "cv_template", "classic") or "classic"),
                cv_color_scheme=str(getattr(args, "cv_color_scheme", "classic_navy") or "classic_navy"),
                languages=list(args.languages),
                profile_image_path=profile_image_path,
                include_profile_image=include_photo,
                profile_links=profile_links,
                output_path=docx_path,
            )
            docx_created = docx_path.exists()
        except Exception as exc:
            docx_error = str(exc)
            docx_created = write_docx(docx_path, cv_text)

        if docx_created and bool(getattr(args, "generate_pdf", True)):
            try:
                generated_pdf_path = Path(convert_docx_to_pdf(str(docx_path)))
                if generated_pdf_path.exists() and generated_pdf_path != pdf_path:
                    generated_pdf_path.replace(pdf_path)
                pdf_created = pdf_path.exists() or generated_pdf_path.exists()
            except Exception as exc:
                pdf_generation_error = str(exc)

        cv_index_records.append(
            {
                "category_id": category_id,
                "category_name": category_name,
                "jobs_assigned_count": int(jobs_per_category.get(category_id, 0)),
                "cv_txt_path": str(txt_path),
                "cv_docx_path": str(docx_path) if docx_created else "",
                "cv_pdf_path": str(pdf_path) if pdf_path.exists() else "",
                "docx_created": bool(docx_created),
                "pdf_created": bool(pdf_path.exists()),
                "docx_error": docx_error,
                "pdf_generation_error": pdf_generation_error,
            }
        )

    index_payload = {"role_cv_count": len(cv_index_records), "role_cvs": cv_index_records}
    index_path = resolve_path(args.role_cv_index_json)
    save_json_file(index_path, index_payload)
    print("Stage 4 complete.")
    print(f"Generated role CVs: {len(cv_index_records)}")
    print(f"Role CV directory: {role_cv_output_dir}")
    print(f"Role CV index: {index_path}")
    return {
        "role_cv_index": index_payload,
        "role_cv_records": cv_index_records,
        "role_cv_output_dir": str(role_cv_output_dir),
        "role_cv_index_path": str(index_path),
    }
