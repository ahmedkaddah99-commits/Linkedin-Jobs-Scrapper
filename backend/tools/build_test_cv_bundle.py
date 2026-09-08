from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from backend.capabilities.reusable_packages.reusable_profiles import (
    _parse_baseline_profile,
    build_role_cv_record,
)
from backend.capabilities.tailored_documents.rendering import (
    CV_COLOR_SCHEMES,
    CV_FONT_OPTIONS,
    CV_TEMPLATE_PRESETS,
    create_cv_document,
    resolve_optional_image_path,
    resolve_profile_image_path,
)
from backend.config.job_seeker import (
    cfg_bool,
    cfg_list,
    cfg_str,
    load_job_seeker_config,
    normalize_windows_env_path,
)
from backend.profiles.reusable_packages import load_baseline_profile_text
from backend.tools.build_test_cv_web_templates import build_templates as build_web_templates


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "test CV"
ROLE_SOURCE_DIR = ROOT / "backend" / "config" / "outputs" / "role_cvs"
ROLE_INDEX_PATH = ROOT / "backend" / "config" / "outputs" / "stage4_role_cvs.json"
TAILORED_SOURCE_DIR = ROOT / "backend" / "config" / "outputs" / "generated_docs" / "2026-04-19"


def _clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def _extract_baseline_field(text: str, prefix: str, fallback: str = "") -> str:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if line.lower().startswith(prefix.lower()):
            return line.split(":", 1)[1].strip() if ":" in line else line.strip()
    return fallback


def _role_category_rows() -> list[dict[str, Any]]:
    if ROLE_INDEX_PATH.exists():
        payload = json.loads(ROLE_INDEX_PATH.read_text(encoding="utf-8"))
        rows = payload.get("role_cvs")
        if isinstance(rows, list):
            return [row for row in rows if isinstance(row, dict)]
    return []


def _copy_current_codebase_outputs(base_dir: Path) -> dict[str, Any]:
    role_target_dir = base_dir / "01_current_codebase_outputs" / "role_based_reusable_cvs"
    tailored_target_dir = base_dir / "01_current_codebase_outputs" / "job_specific_tailored_cvs" / "2026-04-19"
    role_target_dir.mkdir(parents=True, exist_ok=True)
    tailored_target_dir.mkdir(parents=True, exist_ok=True)

    copied_role_files: list[str] = []
    copied_tailored_files: list[str] = []

    for source_path in sorted(ROLE_SOURCE_DIR.glob("*")):
        if not source_path.is_file():
            continue
        if source_path.name.startswith("~$"):
            continue
        target_path = role_target_dir / source_path.name
        shutil.copy2(source_path, target_path)
        copied_role_files.append(target_path.relative_to(base_dir).as_posix())

    for job_dir in sorted(TAILORED_SOURCE_DIR.iterdir(), key=lambda item: item.name.lower()):
        if not job_dir.is_dir():
            continue
        target_job_dir = tailored_target_dir / job_dir.name
        target_job_dir.mkdir(parents=True, exist_ok=True)
        for source_path in sorted(job_dir.glob("*_CV.*")):
            if not source_path.is_file():
                continue
            target_path = target_job_dir / source_path.name
            shutil.copy2(source_path, target_path)
            copied_tailored_files.append(target_path.relative_to(base_dir).as_posix())

    return {
        "role_file_count": len(copied_role_files),
        "tailored_cv_file_count": len(copied_tailored_files),
        "role_files": copied_role_files,
        "tailored_cv_files": copied_tailored_files,
    }


def _resolve_candidate_assets(config: dict) -> tuple[Path | None, list[dict[str, str]]]:
    normalized_profile_image = normalize_windows_env_path(cfg_str(config, ("candidate", "profile_image_path"), ""))
    include_photo = cfg_bool(config, ("candidate", "include_photo"), True)
    profile_image_path = resolve_profile_image_path(normalized_profile_image) if include_photo else None

    profile_links: list[dict[str, str]] = []
    for key, default_text, default_icon in (
        ("linkedin", "LinkedIn", "in"),
        ("github", "GitHub", "GH"),
    ):
        link_url = cfg_str(config, ("candidate", "profile_links", key, "url"), "")
        if not link_url:
            continue
        profile_links.append(
            {
                "icon": cfg_str(config, ("candidate", "profile_links", key, "icon"), default_icon),
                "text": cfg_str(config, ("candidate", "profile_links", key, "text"), default_text),
                "url": link_url,
                "logo_path": resolve_optional_image_path(
                    cfg_str(config, ("candidate", "profile_links", key, "logo_path"), "")
                ),
            }
        )
    return profile_image_path, profile_links


def _representative_role_record() -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    baseline_text = load_baseline_profile_text()
    baseline_profile = _parse_baseline_profile(baseline_text)
    candidate_location = _extract_baseline_field(baseline_text, "Wohnort", "Erlangen / Nuremberg")

    category_rows = _role_category_rows()
    representative_category = next(
        (row for row in category_rows if str(row.get("category_id") or "") == "warehouse_logistics"),
        None,
    )
    if representative_category is None:
        representative_category = {
            "category_id": "warehouse_logistics",
            "category_name": "Warehouse / Logistics",
            "jobs_assigned_count": 0,
        }

    category = {
        "id": str(representative_category.get("category_id") or "warehouse_logistics"),
        "name": str(representative_category.get("category_name") or "Warehouse / Logistics"),
    }
    record = build_role_cv_record(
        candidate_location=candidate_location,
        baseline_profile=baseline_profile,
        category=category,
    )
    return record, category, {
        "candidate_location": candidate_location,
        "baseline_text": baseline_text,
    }


def _generate_builtin_style_matrix(base_dir: Path) -> dict[str, Any]:
    config = load_job_seeker_config()
    candidate_name = cfg_str(config, ("candidate", "name"), "Ahmed Kaddah")
    candidate_email = cfg_str(config, ("candidate", "email"), "ahmed.kaddah@tutamail.com")
    languages = [str(item) for item in cfg_list(config, ("candidate", "languages"), []) if str(item).strip()]
    include_photo = cfg_bool(config, ("candidate", "include_photo"), True)
    profile_image_path, profile_links = _resolve_candidate_assets(config)
    record, category, context = _representative_role_record()

    style_dir = base_dir / "02_builtin_style_matrix"
    docx_dir = style_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)

    source_txt_path = style_dir / "warehouse_logistics_source.txt"
    source_txt_path.write_text(
        "\n".join(
            [
                f"Candidate: {candidate_name}",
                f"Email: {candidate_email}",
                f"Category: {category['name']}",
                f"Location: {context['candidate_location']}",
                "",
                str(record.get('cv_professional_summary') or "").strip(),
                "",
                "Skills:",
                *[f"- {item}" for item in record.get("cv_skills") or []],
            ]
        ).strip(),
        encoding="utf-8",
    )

    manifest_rows: list[dict[str, Any]] = []
    for template_id in CV_TEMPLATE_PRESETS:
        for color_scheme_id in CV_COLOR_SCHEMES:
            for font in CV_FONT_OPTIONS:
                font_id = str(font["id"])
                output_name = f"{template_id}__{color_scheme_id}__{font_id}.docx"
                output_path = docx_dir / output_name
                create_cv_document(
                    record=record,
                    docs_dir=docx_dir,
                    run_date="",
                    candidate_name=candidate_name,
                    candidate_email=candidate_email,
                    cv_font_name=font_id,
                    cv_template_id=template_id,
                    cv_color_scheme=color_scheme_id,
                    languages=languages,
                    profile_image_path=profile_image_path,
                    include_profile_image=include_photo,
                    profile_links=profile_links,
                    output_path=output_path,
                )
                manifest_rows.append(
                    {
                        "path": output_path.relative_to(base_dir).as_posix(),
                        "template": template_id,
                        "color_scheme": color_scheme_id,
                        "font": font_id,
                        "include_photo": include_photo and profile_image_path is not None,
                        "representative_category": category["id"],
                    }
                )

    manifest = {
        "representative_category": category,
        "style_combination_count": len(manifest_rows),
        "template_count": len(CV_TEMPLATE_PRESETS),
        "color_scheme_count": len(CV_COLOR_SCHEMES),
        "font_count": len(CV_FONT_OPTIONS),
        "styles": manifest_rows,
    }
    (style_dir / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def _external_resume_payload() -> dict[str, Any]:
    summary = (
        "Operations and logistics support professional with hands-on experience managing more than "
        "1,600 e-scooters across 7 cities, coordinating 7 transport vans, completing daily readiness "
        "checks, and supporting customer-facing field operations. Recognized for fast response times, "
        "reliable execution, onboarding new team members, and disciplined daily documentation. "
        "Available immediately for warehouse, logistics, operational support, and service roles in "
        "the Erlangen / Nuremberg area."
    )
    return {
        "headline": "Operations and Logistics Support",
        "summary": summary,
        "core_competencies": [
            "Fleet coordination",
            "Warehouse and field operations",
            "Transport van logistics",
            "Daily inspections and readiness checks",
            "Inventory counting and documentation",
            "Customer-facing service",
            "Team onboarding",
            "Shift reliability",
            "Arabic (Native)",
            "English (C1)",
            "German (B1/B2)",
        ],
        "experience": [
            {
                "role": "Logistics and Fleet Operations Associate",
                "company": "Zeus Scooters GmbH",
                "period": "Dec 2023 - Jul 2024",
                "bullets": [
                    "Managed day-to-day support for a fleet of more than 1,600 e-scooters across 7 cities.",
                    "Coordinated logistics using 7 transport vans to keep assets available and operational.",
                    "Completed daily maintenance checks and readiness inspections to support reliable service.",
                    "Trained new team members on operational routines and day-to-day execution.",
                ],
            },
            {
                "role": "Service and Operations Associate",
                "company": "Roxy Mobility GmbH",
                "period": "Dec 2020 - Oct 2024",
                "bullets": [
                    "Handled daily field service and customer support in a fast-moving operating environment.",
                    "Coordinated with partners including Deutsche Bahn on site and logistics-related tasks.",
                    "Earned top recognition for the fastest response time in the city.",
                    "Maintained daily records and inventory tracking for operational equipment.",
                ],
            },
            {
                "role": "Project and Team Assistant",
                "company": "General Administration and Logistics",
                "period": "Nov 2024 - Present",
                "bullets": [
                    "Supported daily data capture and task tracking for ongoing team activities.",
                    "Prepared materials for workshops and team sessions.",
                    "Provided general administrative support to improve workflow consistency.",
                ],
            },
        ],
        "availability": "Available immediately for full-time or mini-job roles.",
        "location": "Erlangen / Nuremberg, Germany",
        "email": "ahmed.kaddah@tutamail.com",
        "name": "Ahmed Kaddah",
    }


def _write_external_resume_docx(output_path: Path, payload: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(payload["name"])
    run.bold = True
    run.font.size = Pt(16)

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run(
        f"{payload['headline']} | {payload['location']} | {payload['email']}"
    ).font.size = Pt(10.5)

    def add_heading(text: str) -> None:
        paragraph = doc.add_paragraph()
        heading_run = paragraph.add_run(text.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(11)

    add_heading("Professional Summary")
    summary_paragraph = doc.add_paragraph(payload["summary"])
    summary_paragraph.paragraph_format.space_after = Pt(4)

    add_heading("Core Competencies")
    competencies = doc.add_paragraph(", ".join(payload["core_competencies"]))
    competencies.paragraph_format.space_after = Pt(4)

    add_heading("Experience")
    for experience in payload["experience"]:
        header = doc.add_paragraph()
        header_run = header.add_run(
            f"{experience['role']} | {experience['company']} | {experience['period']}"
        )
        header_run.bold = True
        for bullet in experience["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

    add_heading("Additional Information")
    doc.add_paragraph(payload["availability"])
    doc.add_paragraph("Languages: Arabic (Native), English (C1), German (B1/B2)")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _write_external_better_alternative(base_dir: Path) -> dict[str, Any]:
    alt_dir = base_dir / "03_external_better_alternative"
    alt_dir.mkdir(parents=True, exist_ok=True)

    payload = _external_resume_payload()

    markdown_text = "\n".join(
        [
            f"# {payload['name']}",
            "",
            f"**{payload['headline']}**",
            "",
            f"{payload['location']} | {payload['email']}",
            "",
            "## Professional Summary",
            "",
            payload["summary"],
            "",
            "## Core Competencies",
            "",
            ", ".join(payload["core_competencies"]),
            "",
            "## Experience",
            "",
            *[
                "\n".join(
                    [
                        f"### {item['role']} | {item['company']} | {item['period']}",
                        *[f"- {bullet}" for bullet in item["bullets"]],
                        "",
                    ]
                ).rstrip()
                for item in payload["experience"]
            ],
            "",
            "## Additional Information",
            "",
            f"- {payload['availability']}",
            "- Languages: Arabic (Native), English (C1), German (B1/B2)",
        ]
    ).strip()

    txt_path = alt_dir / "operations_logistics_hybrid_resume.txt"
    md_path = alt_dir / "operations_logistics_hybrid_resume.md"
    docx_path = alt_dir / "operations_logistics_hybrid_resume.docx"
    notes_path = alt_dir / "why_this_is_better.md"

    txt_path.write_text(markdown_text.replace("# ", "").replace("## ", "").replace("### ", ""), encoding="utf-8")
    md_path.write_text(markdown_text + "\n", encoding="utf-8")
    _write_external_resume_docx(docx_path, payload)

    notes_path.write_text(
        "\n".join(
            [
                "# Why This Is Better",
                "",
                "This comparison version uses guidance from current public career-service resources instead of only the",
                "current repo renderer and prompt outputs.",
                "",
                "Changes made:",
                "",
                "- Uses one clear role headline instead of mixed-language role labels.",
                "- Keeps the summary short, targeted, and specific to operations/logistics roles.",
                "- Uses objective competencies instead of generic soft-skill labels where possible.",
                "- Rewrites bullets with stronger action verbs and preserved metrics already present in the source profile.",
                "- Keeps experience in reverse chronological order and removes extra visual noise such as the profile photo.",
                "",
                "Source links used for this rewrite:",
                "",
                "- Europass CV guidance: https://europass.europa.eu/en/create-europass-cv",
                "- University of Pennsylvania resume guidance: https://careerservices.upenn.edu/channels/resume/",
                "- Yale action-verb guidance: https://ocs.yale.edu/resume-action-verbs/",
                "- University of Pennsylvania accomplishment formula: https://careerservices.upenn.edu/blog/2025/01/23/how-and-why-to-track-accomplishments-at-work-internship/",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    return {
        "files": [
            txt_path.relative_to(base_dir).as_posix(),
            md_path.relative_to(base_dir).as_posix(),
            docx_path.relative_to(base_dir).as_posix(),
            notes_path.relative_to(base_dir).as_posix(),
        ]
    }


def _write_root_readme(base_dir: Path, manifest: dict[str, Any]) -> None:
    readme_path = base_dir / "README.md"
    style_manifest = manifest["builtin_style_matrix"]
    current_outputs = manifest["current_codebase_outputs"]
    web_templates = manifest["web_inspired_templates"]

    readme_path.write_text(
        "\n".join(
            [
                "# test CV",
                "",
                "This folder packages what the repo can already produce now, plus one stronger externally-informed",
                "comparison version.",
                "",
                "## Included",
                "",
                f"- Current reusable role CV files copied from the repo: {current_outputs['role_file_count']}",
                f"- Current job-specific tailored CV files copied from the repo: {current_outputs['tailored_cv_file_count']}",
                (
                    "- Built-in style matrix generated from the current renderer: "
                    f"{style_manifest['style_combination_count']} docx files "
                    f"({style_manifest['template_count']} templates x "
                    f"{style_manifest['color_scheme_count']} color schemes x "
                    f"{style_manifest['font_count']} fonts)"
                ),
                f"- Web-inspired print-ready HTML templates: {web_templates['template_file_count']}",
                "- External better-alternative comparison: 1 content version in txt, md, and docx",
                "",
                "## Notes",
                "",
                "- The current repo already supports many more future job-specific CVs, but those depend on new job input",
                "  and AI generation. This bundle includes all checked-in CV outputs already present in the repo plus a",
                "  full built-in style matrix for one representative role profile.",
                "- The web-inspired templates are implemented as standalone HTML files because that is the lightest path",
                "  to attractive, maintainable, print-ready CV generation in code.",
                "- PDF files are not included here because the checked-in repo state shows PDF export failures when Word or",
                "  LibreOffice conversion is unavailable.",
                "- The external comparison version is intentionally simpler and more ATS-oriented so you can compare content",
                "  structure, not just color or template styling.",
                "",
                "## Structure",
                "",
                "- `01_current_codebase_outputs/`",
                "- `02_builtin_style_matrix/`",
                "- `03_external_better_alternative/`",
                "- `04_web_inspired_templates/`",
                "",
                "## Source Links For The External Comparison",
                "",
                "- https://europass.europa.eu/en/create-europass-cv",
                "- https://careerservices.upenn.edu/channels/resume/",
                "- https://ocs.yale.edu/resume-action-verbs/",
                "- https://careerservices.upenn.edu/blog/2025/01/23/how-and-why-to-track-accomplishments-at-work-internship/",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def build_bundle() -> dict[str, Any]:
    _clean_dir(OUTPUT_DIR)

    current_outputs = _copy_current_codebase_outputs(OUTPUT_DIR)
    builtin_style_matrix = _generate_builtin_style_matrix(OUTPUT_DIR)
    external_better_alternative = _write_external_better_alternative(OUTPUT_DIR)
    web_template_paths = build_web_templates()

    manifest = {
        "output_dir": str(OUTPUT_DIR),
        "current_codebase_outputs": current_outputs,
        "builtin_style_matrix": builtin_style_matrix,
        "external_better_alternative": external_better_alternative,
        "web_inspired_templates": {
            "template_file_count": len(web_template_paths),
            "files": [path.relative_to(OUTPUT_DIR).as_posix() for path in web_template_paths],
        },
    }
    (OUTPUT_DIR / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    _write_root_readme(OUTPUT_DIR, manifest)
    return manifest


def main() -> int:
    manifest = build_bundle()
    print(json.dumps(manifest, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
