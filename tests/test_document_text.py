import io
import unittest
from unittest.mock import patch

import fitz
from docx import Document
from PIL import Image

from backend.profiles.document_text import create_word_companion_bytes, extract_document_text, extraction_metadata


def _pdf_bytes(*, text: str | None = None, page_texts: list[str] | None = None) -> bytes:
    document = fitz.open()
    for page_text in page_texts if page_texts is not None else [text or ""]:
        page = document.new_page()
        if page_text:
            page.insert_text((72, 72), page_text)
    data = document.tobytes()
    document.close()
    return data


class DocumentTextTests(unittest.TestCase):
    def test_extracts_embedded_pdf_text_when_ocr_is_disabled(self):
        with patch("backend.profiles.document_text._ocr_image") as ocr_image:
            extraction = extract_document_text(
                "resume.pdf",
                _pdf_bytes(text="Embedded candidate summary"),
                allow_ocr=False,
            )

        self.assertEqual(extraction["text"], "Embedded candidate summary")
        self.assertEqual(extraction["method"], "pdf_native")
        self.assertEqual(extraction["warnings"], [])
        ocr_image.assert_not_called()

    def test_native_pdf_text_prevents_ocr_when_ocr_is_enabled(self):
        with patch("backend.profiles.document_text._ocr_image") as ocr_image:
            extraction = extract_document_text(
                "resume.pdf",
                _pdf_bytes(text="Native text"),
                allow_ocr=True,
            )

        self.assertEqual(extraction["method"], "pdf_native")
        ocr_image.assert_not_called()

    def test_extracts_plain_text_and_builds_metadata(self):
        extraction = extract_document_text("notes.txt", b"Runr document text")

        self.assertEqual(extraction["text"], "Runr document text")
        self.assertEqual(extraction["method"], "plain_text")
        self.assertEqual(extraction_metadata(extraction)["source_char_count"], 18)

    def test_extracts_docx_tables_and_paragraphs(self):
        document = Document()
        document.add_paragraph("Candidate summary")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Python"
        output = io.BytesIO()
        document.save(output)

        extraction = extract_document_text("resume.docx", output.getvalue())

        self.assertIn("Candidate summary", extraction["text"])
        self.assertIn("Skill | Python", extraction["text"])
        self.assertEqual(extraction["method"], "docx")
        self.assertEqual(extraction["warnings"], [])

    def test_word_companion_contains_extracted_text(self):
        companion = create_word_companion_bytes("First line\nSecond line", title="CV")
        document = Document(io.BytesIO(companion))

        self.assertEqual([paragraph.text for paragraph in document.paragraphs], ["First line", "Second line"])

    def test_unknown_binary_type_reports_warning(self):
        extraction = extract_document_text("archive.bin", b"\x00\x01\x02")

        self.assertEqual(extraction["text"], "")
        self.assertTrue(extraction["warnings"])

    def test_fast_pdf_extraction_skips_ocr_when_native_text_is_missing(self):
        extraction = extract_document_text("scan.pdf", _pdf_bytes(), allow_ocr=False)

        self.assertEqual(extraction["text"], "")
        self.assertEqual(extraction["method"], "pdf_ocr_skipped")
        self.assertEqual(extraction["char_count"], 0)
        self.assertIn("OCR was skipped", extraction["warnings"][0])
        self.assertIn("Scanned PDFs need OCR", extraction["warnings"][0])

    def test_uses_ocr_for_scanned_pdf_when_enabled(self):
        with patch("backend.profiles.document_text._ocr_image", return_value="OCR candidate summary") as ocr_image:
            extraction = extract_document_text("scan.pdf", _pdf_bytes(), allow_ocr=True)

        self.assertEqual(extraction["text"], "OCR candidate summary")
        self.assertEqual(extraction["method"], "pdf_ocr")
        self.assertEqual(extraction["warnings"], [])
        ocr_image.assert_called_once()

    def test_uses_page_level_ocr_for_mixed_pdf(self):
        ocr_page = {
            "text": "OCR second page",
            "page_number": 2,
            "method": "ocr",
            "status": "ready",
            "char_count": 15,
            "confidence": 0.82,
            "rotation_degrees": 0,
            "warnings": [],
        }
        with patch("backend.profiles.document_text._ocr_image_with_metadata", return_value=ocr_page):
            extraction = extract_document_text(
                "mixed.pdf",
                _pdf_bytes(page_texts=["Native first page content", ""]),
                allow_ocr=True,
            )

        self.assertEqual(extraction["method"], "pdf_mixed")
        self.assertEqual([page["method"] for page in extraction["pages"]], ["native", "ocr"])
        self.assertIn("Native first page content", extraction["text"])
        self.assertIn("OCR second page", extraction["text"])

    def test_rotates_low_resolution_image_and_records_confidence(self):
        output = io.BytesIO()
        Image.new("RGB", (400, 300), "white").save(output, format="PNG")
        with (
            patch("backend.profiles.document_text._ocr_image", return_value="Rotated OCR text"),
            patch("pytesseract.image_to_osd", return_value={"rotate": 90}),
            patch("pytesseract.image_to_data", return_value={"conf": ["80"]}),
        ):
            extraction = extract_document_text("scan.png", output.getvalue(), allow_ocr=True)

        self.assertEqual(extraction["pages"][0]["rotation_degrees"], 90)
        self.assertEqual(extraction["confidence"], 0.8)
        self.assertIn("low resolution", " ".join(extraction["warnings"]))
        self.assertEqual(extraction_metadata(extraction)["text_extraction"]["pages"][0]["page_number"], 1)

    def test_malformed_pdf_returns_safe_warning(self):
        document_content = "private candidate content"
        extraction = extract_document_text("broken.pdf", document_content.encode(), allow_ocr=True)

        self.assertEqual(extraction["text"], "")
        self.assertEqual(extraction["method"], "none")
        self.assertTrue(extraction["warnings"])
        self.assertNotIn(document_content, " ".join(extraction["warnings"]))


if __name__ == "__main__":
    unittest.main()
