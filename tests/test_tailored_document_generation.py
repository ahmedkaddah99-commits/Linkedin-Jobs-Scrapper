import subprocess
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from backend.capabilities.tailored_documents.application_requirements import detect_application_requirements
from backend.capabilities.tailored_documents.cv_structuring import ensure_structured_cv_fields
from backend.capabilities.tailored_documents.documents import _resolve_profile_link_url, _stage4_generation_fingerprint
from backend.capabilities.tailored_documents.generation import build_docs_prompt, generate_docs_for_job
from backend.capabilities.tailored_documents.language_rules import detect_reasons
from backend.capabilities.tailored_documents.modes import resolve_cv_generation_prompt_settings
from backend.capabilities.tailored_documents.rendering import (
    build_cv_html_export_payload,
    create_cv_document,
    create_cv_pdf_document,
)
from backend.capabilities.tailored_documents.runtime import build_main_defaults, build_stage4_args


def _draft(summary: str, *, skill: str = "SQL", bullet: str = "Delivered reporting improvements.") -> dict:
    return {
        "cv_professional_summary": summary,
        "cv_professional_experience": [
            {
                "role_title": "Business Analyst",
                "company": "ACME",
                "period": "2022-2024",
                "bullets": [bullet, "Improved stakeholder reporting."],
            }
        ],
        "cv_education": [
            {
                "degree_title": "MSc Information Systems",
                "thesis_title": "Analytics Thesis",
                "thesis_bullets": ["Built dashboards."],
            }
        ],
        "cv_skills": [skill, "Stakeholder Management", "Requirements Gathering"],
        "tailored_cv": f"Professional Summary: {summary}",
    }


class TailoredDocumentGenerationTests(unittest.TestCase):
    def setUp(self):
        self.job = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "full_description": "Need SQL, stakeholder management, and dashboarding.",
        }
        self.cv_text = "\n".join(
            [
                "Professional Summary",
                "Trusted analyst with delivery experience.",
                "",
                "Professional Experience",
                "Business Analyst | ACME | 2022-2024",
                "- Baseline bullet one.",
                "- Baseline bullet two.",
                "",
                "Education",
                "MSc Information Systems",
                "Master Thesis: Analytics Thesis",
                "- Built baseline dashboards.",
                "",
                "Projects",
                "Insight Automation",
                "- Built a baseline workflow.",
            ]
        )

    def test_language_rules_do_not_treat_german_umlauts_as_french(self):
        reasons = detect_reasons(
            {
                "title": "Product Manager",
                "full_description": "Wir suchen Produktarbeit fuer Muenchen mit Verantwortung f\u00fcr Teams.",
            },
            german_special_char_threshold=9999,
            french_special_char_threshold=0,
            spanish_special_char_threshold=9999,
            max_german_level="B2",
            profile_languages=["English - C1", "German - B2"],
        )

        self.assertFalse(any("French" in reason for reason in reasons))

    def test_language_rules_use_configured_language_levels(self):
        reasons = detect_reasons(
            {
                "title": "Product Analyst",
                "full_description": "This role requires fluent English at C1 level.",
            },
            german_special_char_threshold=9999,
            french_special_char_threshold=9999,
            spanish_special_char_threshold=9999,
            max_german_level="B2",
            profile_languages=["Chinese - C1", "English - B2"],
        )

        self.assertIn("English level requirement (C1) is above saved level (B2).", reasons)

    def test_language_rules_allow_required_language_when_profile_matches(self):
        reasons = detect_reasons(
            {
                "title": "Market Analyst",
                "full_description": "Mandarin Chinese proficiency is required for this role.",
            },
            german_special_char_threshold=9999,
            french_special_char_threshold=9999,
            spanish_special_char_threshold=9999,
            max_german_level="B2",
            profile_languages=["Chinese - C1", "English - B2"],
        )

        self.assertEqual(reasons, [])

    def test_application_requirements_plan_translation_for_requested_german_cv(self):
        requirements = detect_application_requirements(
            {
                "title": "Operations Analyst",
                "full_description": "Please submit your CV in German together with a cover letter.",
            },
            cv_text=self.cv_text,
            cv_can_translate=True,
        )

        language = requirements["cv_requirements"]["language"]
        self.assertEqual(language["target_language"], "German")
        self.assertEqual(language["source_language"], "English")
        self.assertEqual(language["output_language"], "German")
        self.assertTrue(language["translation_required"])
        self.assertTrue(language["will_translate"])
        self.assertTrue(
            any(item["code"] == "cv_language_translation_planned" for item in requirements["warnings"])
        )
        self.assertTrue(
            any(item["document_type"] == "motivation_letter" for item in requirements["required_documents"])
        )

    def test_application_requirements_block_standard_cv_language_conflict(self):
        requirements = detect_application_requirements(
            {
                "title": "Operations Analyst",
                "full_description": "Bitte reichen Sie Ihren Lebenslauf auf Deutsch ein.",
            },
            cv_text=self.cv_text,
            cv_can_translate=False,
        )

        self.assertTrue(
            any(
                item["code"] == "cv_language_conflict" and item["severity"] == "blocking"
                for item in requirements["warnings"]
            )
        )

    def test_application_requirements_do_not_mistake_language_skill_for_cv_language(self):
        requirements = detect_application_requirements(
            {
                "title": "Operations Analyst",
                "full_description": "German B2 is required. Please submit your CV and certificates online.",
            },
            cv_text=self.cv_text,
            cv_can_translate=True,
        )

        self.assertEqual(requirements["cv_requirements"]["language"]["target_language"], "")

    def test_build_docs_prompt_uses_selected_output_language(self):
        prompt = build_docs_prompt(
            cv_text=self.cv_text,
            job={**self.job, "full_description": "Bitte reichen Sie Ihren Lebenslauf auf Deutsch ein."},
            candidate_name="Ahmed",
            cv_generation_mode="aggressive_customization",
            output_language="German",
        )

        self.assertIn("Write tailored CV content in German.", prompt)
        self.assertIn("Keep JSON keys exactly as specified in English.", prompt)
        self.assertNotIn("Write tailored CV content in English.", prompt)

    @patch("backend.capabilities.tailored_documents.generation._improve_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._score_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._generate_structured_cv_once")
    def test_generate_docs_for_job_passes_after_first_score(
        self,
        generate_mock,
        score_mock,
        improve_mock,
    ):
        generate_mock.return_value = _draft("Strong first draft.")
        score_mock.return_value = {
            "score": 93,
            "missing_requirements": [],
            "improvement_actions": [],
            "rationale": "Strong fit.",
        }

        result = generate_docs_for_job(
            deepseek_api_key=None,
            deepseek_model="deepseek-chat",
            gemini_client=None,
            gemini_model="gemini-2.5-flash",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            cv_generation_mode="aggressive_customization",
            extra_instructions="",
            prompt_override="",
            retries=1,
            retry_sleep=0.0,
        )

        self.assertEqual(result["ats_attempt_count"], 1)
        self.assertEqual(result["ats_score"], 93)
        self.assertEqual(result["ats_gate_state"], "passed")
        self.assertTrue(result["ats_can_export_final"])
        self.assertFalse(result["ats_export_anyway_allowed"])
        self.assertEqual(result["ats_export_gate"]["gate_state"], "passed")
        improve_mock.assert_not_called()

    @patch("backend.capabilities.tailored_documents.generation._improve_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._score_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._generate_structured_cv_once")
    def test_generate_docs_for_job_stops_when_score_stalls_and_keeps_best_attempt(
        self,
        generate_mock,
        score_mock,
        improve_mock,
    ):
        first_draft = _draft("First ATS draft.", skill="SQL")
        second_draft = _draft("Second ATS draft.", skill="Python", bullet="Refined the wording.")
        generate_mock.return_value = first_draft
        improve_mock.return_value = second_draft
        score_mock.side_effect = [
            {
                "score": 84,
                "missing_requirements": ["SQL"],
                "improvement_actions": ["Make SQL evidence more explicit."],
                "rationale": "Close match.",
            },
            {
                "score": 84,
                "missing_requirements": ["Python"],
                "improvement_actions": ["Add more evidence."],
                "rationale": "No improvement.",
            },
        ]

        result = generate_docs_for_job(
            deepseek_api_key=None,
            deepseek_model="deepseek-chat",
            gemini_client=None,
            gemini_model="gemini-2.5-flash",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            cv_generation_mode="aggressive_customization",
            extra_instructions="",
            prompt_override="",
            retries=1,
            retry_sleep=0.0,
        )

        self.assertEqual(result["cv_professional_summary"], first_draft["cv_professional_summary"])
        self.assertEqual(result["ats_score"], 84)
        self.assertEqual(result["ats_attempt_count"], 2)
        self.assertEqual(result["ats_stop_reason"], "score_stalled")
        self.assertEqual(result["ats_missing_requirements"], ["SQL"])
        self.assertEqual(result["ats_gate_state"], "blocked")
        self.assertFalse(result["ats_can_export_final"])
        self.assertTrue(result["ats_export_anyway_allowed"])
        self.assertEqual(result["ats_export_gate"]["metadata"]["stop_reason"], "score_stalled")
        self.assertIn("Best score reached: 84%", result["ats_last_warning"])
        self.assertEqual(len(result["ats_attempt_history"]), 2)
        self.assertEqual(result["ats_attempt_history"][0]["changed_sections"], ["initial_draft"])
        self.assertEqual(result["ats_attempt_history"][1]["changed_sections"], ["summary", "experience", "skills"])
        self.assertEqual(improve_mock.call_count, 1)

    @patch("backend.capabilities.tailored_documents.generation._improve_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._score_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._generate_structured_cv_once")
    def test_generate_docs_for_job_blocks_after_third_scored_attempt(
        self,
        generate_mock,
        score_mock,
        improve_mock,
    ):
        generate_mock.return_value = _draft("Draft one.", skill="Excel")
        improve_mock.side_effect = [
            _draft("Draft two.", skill="SQL"),
            _draft("Draft three.", skill="Dashboarding"),
        ]
        score_mock.side_effect = [
            {
                "score": 80,
                "missing_requirements": ["SQL"],
                "improvement_actions": ["Add SQL evidence."],
                "rationale": "Weak SQL coverage.",
            },
            {
                "score": 85,
                "missing_requirements": ["Dashboarding"],
                "improvement_actions": ["Highlight dashboard work."],
                "rationale": "Improved.",
            },
            {
                "score": 88,
                "missing_requirements": ["Leadership"],
                "improvement_actions": ["Clarify leadership scope."],
                "rationale": "Best attempt but under target.",
            },
        ]

        result = generate_docs_for_job(
            deepseek_api_key=None,
            deepseek_model="deepseek-chat",
            gemini_client=None,
            gemini_model="gemini-2.5-flash",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            cv_generation_mode="aggressive_customization",
            extra_instructions="",
            prompt_override="",
            retries=1,
            retry_sleep=0.0,
        )

        self.assertEqual(result["cv_professional_summary"], "Draft three.")
        self.assertEqual(result["ats_score"], 88)
        self.assertEqual(result["ats_attempt_count"], 3)
        self.assertEqual(result["ats_stop_reason"], "max_attempts_reached")
        self.assertEqual(result["ats_gate_state"], "blocked")
        self.assertTrue(result["ats_export_anyway_allowed"])
        self.assertEqual(result["ats_export_gate"]["best_score"], 88)
        self.assertIn("Best score reached: 88%", result["ats_export_gate"]["last_warning"])
        self.assertEqual(
            [entry["change_summary"] for entry in result["ats_attempt_history"]],
            [
                "Initial tailored CV draft scored.",
                "Updated summary, skills since the previous scored pass.",
                "Updated summary, skills since the previous scored pass.",
            ],
        )
        self.assertEqual(improve_mock.call_count, 2)

    def test_resolve_cv_generation_prompt_settings_uses_mode_specific_fields(self):
        settings = SimpleNamespace(
            light_customization_extra_prompt="Light extra",
            light_customization_prompt_override="Light override",
            aggressive_customization_extra_prompt="Aggressive extra",
            aggressive_customization_prompt_override="Aggressive override",
            stage4_extra_prompt="Legacy extra",
            stage4_prompt_override="Legacy override",
        )

        self.assertEqual(
            resolve_cv_generation_prompt_settings("light_customization", settings),
            ("Light extra", "Light override"),
        )
        self.assertEqual(
            resolve_cv_generation_prompt_settings("aggressive_customization", settings),
            ("Aggressive extra", "Aggressive override"),
        )
        self.assertEqual(
            resolve_cv_generation_prompt_settings(
                "aggressive_customization",
                SimpleNamespace(
                    aggressive_customization_extra_prompt="",
                    aggressive_customization_prompt_override="",
                    stage4_extra_prompt="Legacy extra",
                    stage4_prompt_override="Legacy override",
                ),
            ),
            ("Legacy extra", "Legacy override"),
        )

    def test_stage4_generation_fingerprint_changes_when_export_style_changes(self):
        base = dict(
            cv_generation_mode="aggressive_customization",
            extra_prompt="",
            prompt_override="",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            candidate_email="ahmed@example.com",
            cv_font_name="Calibri",
            cv_template_id="classic",
            cv_color_scheme="classic_navy",
            include_photo=True,
            languages=["English - C1"],
            profile_image_path="user_config/photo.png",
        )

        original = _stage4_generation_fingerprint(**base)
        styled = _stage4_generation_fingerprint(
            **{
                **base,
                "cv_template_id": "modern",
                "cv_color_scheme": "forest",
                "cv_font_name": "Georgia",
                "include_photo": False,
            }
        )

        self.assertNotEqual(original, styled)

    def test_stage4_generation_fingerprint_changes_when_job_description_changes(self):
        base = dict(
            cv_generation_mode="aggressive_customization",
            extra_prompt="",
            prompt_override="",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            candidate_email="ahmed@example.com",
            cv_font_name="Calibri",
            cv_template_id="classic",
            cv_color_scheme="classic_navy",
            include_photo=True,
            languages=["English - C1"],
            profile_image_path="user_config/photo.png",
        )

        original = _stage4_generation_fingerprint(**base)
        changed_job = {**self.job, "full_description": "Need SQL, product analytics, and German B2."}
        changed = _stage4_generation_fingerprint(**{**base, "job": changed_job})

        self.assertNotEqual(original, changed)

    def test_stage4_generation_fingerprint_changes_when_output_language_changes(self):
        base = dict(
            cv_generation_mode="aggressive_customization",
            extra_prompt="",
            prompt_override="",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            candidate_email="ahmed@example.com",
            cv_font_name="Calibri",
            cv_template_id="classic",
            cv_color_scheme="classic_navy",
            include_photo=True,
            languages=["English - C1"],
            profile_image_path="user_config/photo.png",
        )

        english = _stage4_generation_fingerprint(**{**base, "cv_output_language": "English"})
        german = _stage4_generation_fingerprint(**{**base, "cv_output_language": "German"})

        self.assertNotEqual(english, german)

    @patch("backend.capabilities.tailored_documents.generation._score_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._generate_structured_cv_once")
    def test_light_mode_clamps_forbidden_changes_after_generation(self, generate_mock, score_mock):
        generate_mock.return_value = {
            "cv_professional_summary": "Light-mode summary tuned to the role.",
            "cv_professional_experience": [
                {
                    "role_title": "Senior Analyst",
                    "company": "Different Company",
                    "period": "2020-2021",
                    "bullets": ["Rewritten forbidden bullet."],
                }
            ],
            "cv_education": [
                {
                    "degree_title": "Renamed Degree",
                    "thesis_title": "Renamed Thesis",
                    "thesis_bullets": ["Forbidden education rewrite."],
                }
            ],
            "cv_skills": ["SQL", "Dashboarding", "Stakeholder Management"],
            "tailored_cv": "Light-mode summary tuned to the role.",
        }
        score_mock.return_value = {
            "score": 92,
            "missing_requirements": [],
            "improvement_actions": [],
            "rationale": "Looks good.",
        }

        result = generate_docs_for_job(
            deepseek_api_key=None,
            deepseek_model="deepseek-chat",
            gemini_client=None,
            gemini_model="gemini-2.5-flash",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            cv_generation_mode="light_customization",
            extra_instructions="",
            prompt_override="",
            retries=1,
            retry_sleep=0.0,
            payload_postprocessor=lambda payload: self._clamp_payload(payload, "light_customization"),
        )

        self.assertEqual(result["cv_professional_summary"], "Light-mode summary tuned to the role.")
        self.assertEqual(result["cv_skills"], ["SQL", "Dashboarding", "Stakeholder Management"])
        self.assertEqual(result["cv_professional_experience"][0]["role_title"], "Business Analyst")
        self.assertEqual(result["cv_professional_experience"][0]["company"], "ACME")
        self.assertEqual(result["cv_professional_experience"][0]["period"], "2022-2024")
        self.assertEqual(
            result["cv_professional_experience"][0]["bullets"],
            ["Baseline bullet one.", "Baseline bullet two."],
        )
        self.assertEqual(result["cv_education"][0]["degree_title"], "MSc Information Systems")
        self.assertEqual(result["cv_education"][0]["thesis_bullets"], ["Built baseline dashboards."])

    def test_light_mode_keeps_generated_translation_when_cv_language_changes(self):
        record = {
            "cv_output_language": "German",
            "application_requirements": {
                "cv_requirements": {
                    "language": {
                        "source_language": "English",
                        "output_language": "German",
                        "target_language": "German",
                        "will_translate": True,
                    }
                }
            },
            "cv_professional_summary": "Deutsch zugeschnittenes Profil.",
            "cv_professional_experience": [
                {
                    "role_title": "Business Analyst",
                    "company": "ACME",
                    "period": "2022-2024",
                    "bullets": ["Uebersetzte Basisleistung.", "Verbesserte Stakeholder-Berichte."],
                }
            ],
            "cv_education": [
                {
                    "degree_title": "MSc Information Systems",
                    "thesis_title": "Analytics Thesis",
                    "thesis_bullets": ["Erstellte Dashboards."],
                }
            ],
            "cv_skills": ["SQL"],
        }

        ensure_structured_cv_fields(
            record,
            candidate_name="Ahmed",
            cv_text=self.cv_text,
            cv_generation_mode="light_customization",
        )

        self.assertEqual(
            record["cv_professional_experience"][0]["bullets"],
            ["Uebersetzte Basisleistung.", "Verbesserte Stakeholder-Berichte."],
        )
        self.assertEqual(record["cv_education"][0]["thesis_bullets"], ["Erstellte Dashboards."])

    @patch("backend.capabilities.tailored_documents.generation._score_structured_cv_once")
    @patch("backend.capabilities.tailored_documents.generation._generate_structured_cv_once")
    def test_aggressive_mode_keeps_identity_but_allows_bullet_rewrites(self, generate_mock, score_mock):
        generate_mock.return_value = {
            "cv_professional_summary": "Aggressive summary tuned to the role.",
            "cv_professional_experience": [
                {
                    "role_title": "Business Analyst",
                    "company": "ACME",
                    "period": "2022-2024",
                    "bullets": [
                        "Rewritten aggressive bullet one.",
                        "Rewritten aggressive bullet two.",
                        "Extra bullet that should be discarded.",
                    ],
                }
            ],
            "cv_education": [
                {
                    "degree_title": "Renamed Degree",
                    "thesis_title": "Renamed Thesis",
                    "thesis_bullets": ["Forbidden education rewrite."],
                }
            ],
            "cv_skills": ["SQL", "Dashboarding", "Stakeholder Management"],
            "tailored_cv": "Aggressive summary tuned to the role.",
        }
        score_mock.return_value = {
            "score": 91,
            "missing_requirements": [],
            "improvement_actions": [],
            "rationale": "Looks good.",
        }

        result = generate_docs_for_job(
            deepseek_api_key=None,
            deepseek_model="deepseek-chat",
            gemini_client=None,
            gemini_model="gemini-2.5-flash",
            cv_text=self.cv_text,
            job=self.job,
            candidate_name="Ahmed",
            cv_generation_mode="aggressive_customization",
            extra_instructions="",
            prompt_override="",
            retries=1,
            retry_sleep=0.0,
            payload_postprocessor=lambda payload: self._clamp_payload(payload, "aggressive_customization"),
        )

        self.assertEqual(
            result["cv_professional_experience"][0]["bullets"],
            ["Rewritten aggressive bullet one.", "Rewritten aggressive bullet two."],
        )
        self.assertEqual(result["cv_professional_experience"][0]["role_title"], "Business Analyst")
        self.assertEqual(result["cv_education"][0]["degree_title"], "MSc Information Systems")
        self.assertEqual(result["cv_education"][0]["thesis_bullets"], ["Built baseline dashboards."])

    def test_extracts_baseline_experience_from_plain_experience_heading(self):
        record = {
            "title": "Product Analyst",
            "company": "Example",
            "cv_professional_summary": "Tailored summary.",
            "cv_professional_experience": [],
            "cv_skills": [],
        }
        cv_text = "\n".join(
            [
                "Summary",
                "Experienced analyst.",
                "Experience",
                "Business Analyst | Example GmbH",
                "2022 - Present",
                "- Built dashboard reporting",
                "- Improved fulfillment workflow",
                "Skills",
                "SQL",
            ]
        )

        ensure_structured_cv_fields(
            record,
            candidate_name="Ahmed",
            cv_text=cv_text,
            cv_generation_mode="aggressive_customization",
        )

        self.assertEqual(record["cv_professional_experience"][0]["role_title"], "Business Analyst")
        self.assertEqual(record["cv_professional_experience"][0]["company"], "Example GmbH")
        self.assertEqual(record["cv_professional_experience"][0]["period"], "2022 - Present")
        self.assertEqual(
            record["cv_professional_experience"][0]["bullets"],
            ["Built dashboard reporting", "Improved fulfillment workflow"],
        )

    def test_recovers_tailored_cv_bullets_when_structured_experience_is_empty(self):
        record = {
            "title": "Product Support Specialist",
            "company": "Bruker",
            "cv_professional_summary": "Tailored summary.",
            "cv_professional_experience": [
                {
                    "role_title": "Allianz Technology",
                    "company": "Germany",
                    "period": "May 2025 - Sep 2025",
                    "bullets": [],
                }
            ],
            "cv_skills": ["Product Support"],
            "tailored_cv": "\n".join(
                [
                    "Professional Experience:",
                    "Technology Transformation & Strategy Consulting (Internship) | Allianz Technology | May 2025 - Sep 2025",
                    "- Supported enterprise AI transformation programs for product teams.",
                    "- Designed AI-enabled workflows for support analytics.",
                    "Skills:",
                    "- Product Support",
                ]
            ),
        }
        cv_text = "\n".join(
            [
                "Professional Experience",
                "Allianz Technology | Germany | May 2025 - Sep 2025",
            ]
        )

        ensure_structured_cv_fields(
            record,
            candidate_name="Ahmed",
            cv_text=cv_text,
            cv_generation_mode="aggressive_customization",
        )

        self.assertEqual(
            record["cv_professional_experience"][0]["role_title"],
            "Technology Transformation & Strategy Consulting (Internship)",
        )
        self.assertEqual(record["cv_professional_experience"][0]["company"], "Allianz Technology")
        self.assertEqual(
            record["cv_professional_experience"][0]["bullets"],
            [
                "Supported enterprise AI transformation programs for product teams.",
                "Designed AI-enabled workflows for support analytics.",
            ],
        )

    def test_create_cv_document_writes_visible_experience_bullet_prefixes(self):
        record = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "location_raw": "Berlin, Germany",
            "cv_professional_summary": "Tailored summary.",
            "cv_professional_experience": [
                {
                    "role_title": "Business Analyst",
                    "company": "ACME",
                    "period": "2022-2024",
                    "bullets": ["Delivered reporting improvements."],
                }
            ],
            "cv_skills": ["SQL"],
            "cv_education": [],
        }

        with TemporaryDirectory() as temp_dir:
            output_path = create_cv_document(
                record,
                docs_dir=Path(temp_dir),
                run_date="2026-05-28",
                candidate_name="Ahmed",
                candidate_email="ahmed@example.com",
                cv_font_name="Calibri",
                cv_template_id="classic",
                cv_color_scheme="classic_navy",
                languages=[],
                profile_image_path=None,
                include_profile_image=False,
                profile_links=[],
            )
            exported = Document(output_path)
            texts = [paragraph.text.strip() for paragraph in exported.paragraphs if paragraph.text.strip()]

        self.assertIn("\u2022 Delivered reporting improvements.", texts)

    def test_create_plain_document_uses_reference_layout_with_selected_design_tokens(self):
        record = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "location_raw": "Buffalo, New York",
            "cv_professional_summary": "Friendly and engaging team leader.",
            "cv_professional_experience": [
                {
                    "role_title": "Restaurant Manager",
                    "company": "Contoso Bar and Grill",
                    "period": "2022-Present",
                    "bullets": ["Improved customer satisfaction."],
                }
            ],
            "cv_skills": ["Budgeting", "POS systems", "Communication", "Team leadership"],
            "cv_education": [
                {
                    "degree_title": "B.S. in Business Administration | Bigtown College",
                    "thesis_title": "",
                    "thesis_bullets": [],
                }
            ],
        }

        with TemporaryDirectory() as temp_dir:
            output_path = create_cv_document(
                record,
                docs_dir=Path(temp_dir),
                run_date="2026-06-05",
                candidate_name="May Riley",
                candidate_email="m.riley@example.com",
                cv_font_name="Georgia",
                cv_template_id="plain",
                cv_color_scheme="burgundy",
                languages=[],
                profile_image_path=Path(temp_dir) / "unused.png",
                include_profile_image=True,
                profile_links=[
                    {"text": "LinkedIn", "url": "https://linkedin.example/may"},
                    {"text": "GitHub", "url": "https://github.example/may"},
                ],
            )
            exported = Document(output_path)

        section = exported.sections[0]
        paragraph_texts = [paragraph.text.strip() for paragraph in exported.paragraphs if paragraph.text.strip()]
        header_cell = exported.tables[0].cell(0, 0)
        header_run = header_cell.paragraphs[0].runs[0]

        self.assertAlmostEqual(section.page_width.inches, 8.5, places=2)
        self.assertAlmostEqual(section.page_height.inches, 11.0, places=2)
        self.assertAlmostEqual(section.left_margin.inches, 0.8, places=2)
        self.assertEqual(header_cell.text.strip(), "May Riley")
        self.assertEqual(header_run.font.size.pt, 28.0)
        self.assertEqual(header_run.font.name, "Georgia")
        self.assertEqual(str(header_run.font.color.rgb), "7C2D12")
        self.assertIn('w:color="EA580C"', header_cell._tc.xml)
        self.assertIn("Profile", paragraph_texts)
        self.assertIn("Experience", paragraph_texts)
        self.assertIn("Education", paragraph_texts)
        self.assertIn("Skills & Abilities", paragraph_texts)
        self.assertIn("RESTAURANT MANAGER | CONTOSO BAR AND GRILL | 2022-PRESENT", paragraph_texts)
        self.assertIn("Budgeting | POS systems | Communication | Team leadership", paragraph_texts)
        hyperlink_targets = {
            relationship.target_ref
            for relationship in exported.part.rels.values()
            if relationship.reltype.endswith("/hyperlink")
        }
        self.assertIn("https://linkedin.example/may", hyperlink_targets)
        self.assertIn("https://github.example/may", hyperlink_targets)
        self.assertFalse(exported.inline_shapes)

    def test_legacy_teal_resume_id_resolves_to_plain_layout(self):
        record = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "cv_professional_summary": "Tailored summary.",
            "cv_professional_experience": [],
            "cv_skills": [],
            "cv_education": [],
        }

        with TemporaryDirectory() as temp_dir:
            output_path = create_cv_document(
                record,
                docs_dir=Path(temp_dir),
                run_date="2026-06-05",
                candidate_name="May Riley",
                candidate_email="m.riley@example.com",
                cv_font_name="Calibri",
                cv_template_id="teal_resume",
                cv_color_scheme="ocean_teal",
                languages=[],
                profile_image_path=None,
                include_profile_image=False,
                profile_links=[],
            )
            exported = Document(output_path)

        self.assertAlmostEqual(exported.sections[0].page_width.inches, 8.5, places=2)
        self.assertEqual(exported.tables[0].cell(0, 0).text.strip(), "May Riley")

    def test_create_cv_document_localizes_section_labels_for_german_output(self):
        record = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "location_raw": "Berlin, Germany",
            "cv_output_language": "German",
            "cv_professional_summary": "Zielgerichtetes Profil.",
            "cv_professional_experience": [
                {
                    "role_title": "Business Analyst",
                    "company": "ACME",
                    "period": "2022-2024",
                    "bullets": ["Verbesserte Reporting-Prozesse."],
                }
            ],
            "cv_skills": ["SQL"],
            "cv_education": [
                {
                    "degree_title": "MSc Information Systems",
                    "thesis_title": "",
                    "thesis_bullets": [],
                }
            ],
        }

        with TemporaryDirectory() as temp_dir:
            output_path = create_cv_document(
                record,
                docs_dir=Path(temp_dir),
                run_date="2026-05-31",
                candidate_name="Ahmed",
                candidate_email="ahmed@example.com",
                cv_font_name="Calibri",
                cv_template_id="classic",
                cv_color_scheme="classic_navy",
                languages=["Deutsch - B2"],
                profile_image_path=None,
                include_profile_image=False,
                profile_links=[],
            )
            exported = Document(output_path)
            texts = [paragraph.text.strip() for paragraph in exported.paragraphs if paragraph.text.strip()]

        self.assertIn("Profil", texts)
        self.assertIn("Berufserfahrung", texts)
        self.assertIn("Kompetenzen", texts)
        self.assertIn("Ausbildung", texts)

    def test_html_pdf_payload_uses_same_structured_cv_content_as_docx(self):
        record = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "location_raw": "Berlin, Germany",
            "cv_professional_summary": "Tailored summary.",
            "cv_professional_experience": [
                {
                    "role_title": "Business Analyst",
                    "company": "ACME",
                    "period": "2022-2024",
                    "bullets": ["Delivered reporting improvements."],
                }
            ],
            "cv_strategic_initiatives": [
                {
                    "title": "Analytics Automation",
                    "bullets": ["Built reusable reporting workflows."],
                }
            ],
            "cv_skills": ["SQL"],
            "cv_education": [
                {
                    "degree_title": "MSc Information Systems",
                    "thesis_title": "Analytics Thesis",
                    "thesis_bullets": ["Built dashboards."],
                }
            ],
        }

        payload = build_cv_html_export_payload(
            record,
            candidate_name="Ahmed",
            candidate_email="ahmed@example.com",
            cv_font_name="Calibri",
            cv_template_id="modern",
            cv_color_scheme="forest",
            languages=["English - C1"],
            profile_image_path=None,
            include_profile_image=False,
            profile_links=[{"text": "LinkedIn", "url": "https://linkedin.example/ahmed"}],
        )

        self.assertEqual(payload["documents"]["cv_template"], "plain")
        self.assertEqual(payload["documents"]["cv_output_language"], "English")
        self.assertEqual(payload["profile"]["summary"], "Tailored summary.")
        self.assertEqual(payload["profile"]["recent_experience"][0]["bullets"], ["Delivered reporting improvements."])
        self.assertEqual(payload["profile"]["projects"][0]["bullets"], ["Built reusable reporting workflows."])
        self.assertEqual(payload["profile"]["education"][0]["thesis_title"], "Analytics Thesis")
        self.assertEqual(payload["profile"]["linkedin_url"], "https://linkedin.example/ahmed")

    def test_stage4_profile_link_overrides_replace_config_urls(self):
        args = SimpleNamespace(
            linkedin_url="https://linkedin.example/workspace",
            github_url="https://github.example/workspace",
        )
        config = {
            "candidate": {
                "profile_links": {
                    "linkedin": {"url": "https://linkedin.example/config"},
                    "github": {"url": "https://github.example/config"},
                }
            }
        }

        self.assertEqual(
            _resolve_profile_link_url(args, config, "linkedin"),
            "https://linkedin.example/workspace",
        )
        self.assertEqual(
            _resolve_profile_link_url(args, config, "github"),
            "https://github.example/workspace",
        )
        runtime_defaults = build_main_defaults(config)
        runtime_defaults.update(
            linkedin_url=args.linkedin_url,
            github_url=args.github_url,
        )
        stage4_args = build_stage4_args(SimpleNamespace(**runtime_defaults))
        self.assertEqual(stage4_args.linkedin_url, "https://linkedin.example/workspace")
        self.assertEqual(stage4_args.github_url, "https://github.example/workspace")

    def test_create_cv_pdf_document_invokes_shared_browser_renderer(self):
        record = {
            "job_id": "job_1",
            "title": "Senior Analyst",
            "company": "ACME",
            "cv_professional_summary": "Tailored summary.",
            "cv_professional_experience": [
                {
                    "role_title": "Business Analyst",
                    "company": "ACME",
                    "period": "2022-2024",
                    "bullets": ["Delivered reporting improvements."],
                }
            ],
            "cv_skills": ["SQL"],
            "cv_education": [],
        }

        captured_payloads = []

        def fake_run(command, **kwargs):
            output_path = Path(command[command.index("--output") + 1])
            output_path.write_bytes(b"%PDF-1.4\n")
            captured_payloads.append(kwargs["input"])
            return subprocess.CompletedProcess(command, 0, stdout="", stderr="")

        with TemporaryDirectory() as temp_dir, patch(
            "backend.capabilities.tailored_documents.rendering.subprocess.run",
            side_effect=fake_run,
        ):
            output_path = Path(temp_dir) / "cv.pdf"
            rendered_path = create_cv_pdf_document(
                record,
                output_path=output_path,
                candidate_name="Ahmed",
                candidate_email="ahmed@example.com",
                cv_font_name="Calibri",
                cv_template_id="classic",
                cv_color_scheme="classic_navy",
                languages=[],
                profile_image_path=None,
                include_profile_image=False,
                profile_links=[],
            )

        self.assertEqual(Path(rendered_path), output_path.resolve())
        self.assertIn("Delivered reporting improvements.", captured_payloads[0])

    def _clamp_payload(self, payload, mode):
        normalized_payload = dict(payload)
        ensure_structured_cv_fields(
            normalized_payload,
            candidate_name="Ahmed",
            cv_text=self.cv_text,
            cv_generation_mode=mode,
        )
        return normalized_payload


if __name__ == "__main__":
    unittest.main()
