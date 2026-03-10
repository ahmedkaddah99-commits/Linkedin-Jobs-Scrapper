import argparse
import json
import os
import re
import shutil
import subprocess
import time
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from dotenv import load_dotenv
from google import genai
import requests

from cv_profile import load_cv_text
from job_seeker_config import (
    cfg_bool,
    cfg_float,
    cfg_int,
    cfg_list,
    cfg_str,
    load_job_seeker_config,
    normalize_windows_env_path,
)


DEFAULT_CANDIDATE_NAME = "Kaddah Ahmed"
DEFAULT_CANDIDATE_EMAIL = "ahmed.kaddah@tutamail.com"
DEFAULT_CV_FONT = "Calibri"
DEFAULT_LANGUAGES = [
    "Arabic \u2014 Native", "English \u2014 C1", "German \u2014 B1/B2",
]


def load_json_file(path: Path):
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def save_json_file(path: Path, payload) -> None:
    with path.open("w", encoding="utf-8") as file:
        json.dump(payload, file, indent=4, ensure_ascii=False)


def strip_json_fences(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.removeprefix("```json").removeprefix("```").strip()
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3].strip()
    return cleaned


def sanitize_filename(value: str, max_length: int = 90) -> str:
    cleaned = re.sub(r"[^\w\-\. ]+", "", value or "").strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    if not cleaned:
        cleaned = "item"
    return cleaned[:max_length]


def compact_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def extract_city_from_job(job: Dict) -> str:
    location_raw = compact_whitespace(str(job.get("location_raw") or ""))
    if location_raw:
        first_chunk = location_raw.split(",")[0].strip()
        if first_chunk:
            return first_chunk

    description = str(job.get("full_description") or "")
    patterns = [
        r"(?i)\blocation:\s*([A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\- ]+),\s*(?:[A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\- ]+\s*-\s*)?Germany\b",
        r"(?i)\bbased in\s+([A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\- ]+)\b",
        r"(?i)\bin\s+([A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\- ]+),\s*Germany\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, description)
        if match:
            city = compact_whitespace(match.group(1))
            if city:
                return city

    return "Germany"


def format_header_location(job: Dict) -> str:
    city = extract_city_from_job(job)
    if not city or city.strip().lower() == "germany":
        return "Germany"
    city_clean = city.strip()
    if city_clean.lower().endswith(", germany"):
        return city_clean
    return f"{city_clean}, Germany"


def build_docs_prompt(
    cv_text: str,
    job: Dict,
    candidate_name: str,
    extra_instructions: str = "",
    prompt_override: str = "",
) -> str:
    description_excerpt = str(job.get("full_description") or "")[:6000]
    job_id = str(job.get("job_id", ""))
    title = str(job.get("title", ""))
    company = str(job.get("company", ""))
    job_city = extract_city_from_job(job)

    if prompt_override.strip():
        prompt = (
            prompt_override.strip()
            .replace("{{CV_TEXT}}", cv_text)
            .replace("{{JOB_ID}}", job_id)
            .replace("{{JOB_TITLE}}", title)
            .replace("{{JOB_COMPANY}}", company)
            .replace("{{JOB_CITY}}", job_city)
            .replace("{{JOB_DESCRIPTION}}", description_excerpt)
            .replace("{{CANDIDATE_NAME}}", candidate_name)
        )
    else:
        prompt = f"""
You are an expert career writing assistant.

Candidate full CV:
{cv_text}

Job:
- id: {job_id}
- title: {title}
- company: {company}
- city: {job_city}
- description: {description_excerpt}

Write tailored CV content in English.

Output requirements:
- Return only raw JSON (no markdown)
- Keep schema exactly:
{{
  "professional_summary": "text",
  "professional_experience": [
    {{
      "role_title": "text",
      "company": "text",
      "period": "text",
      "bullets": ["text", "text"]
    }}
  ],
  "education": [
    {{
      "degree_title": "text",
      "thesis_title": "text",
      "thesis_bullets": ["text", "text"]
    }}
  ],
  "skills": ["text", "text"]
}}

Content rules:
- professional_summary:
  - detailed and substantial (about 70-140 words)
  - role-specific and company-specific
  - naturally professional, non-generic
  - clearly explain why the candidate is a strong fit
- professional_experience:
  - use only roles that exist in the candidate CV
  - keep role titles exactly as written in the candidate CV
  - preserve the same role order as in the candidate CV
  - 3 to 6 roles total
  - each role has 3 to 5 concise achievement-oriented bullets
  - tailor bullets to the target role and ATS keywords
  - do not fabricate companies or dates
  - if Allianz Technology appears in the CV, include it explicitly
  - never include projects/initiatives in this section
- skills:
  - 16 to 25 ATS-friendly skills relevant to the job description
  - include realistic adjacent skills when useful
- if referencing EDUCATION content from the candidate CV:
  - keep degree/thesis titles exactly as written in the CV
  - do not rename or paraphrase degree/thesis titles
  - preserve the same education item order as in the candidate CV
  - you may only reword thesis bullet wording to fit the target role
  - if referencing PROJECTS from the candidate CV:
  - keep project titles exactly as written in the CV
  - preserve the same project order as in the candidate CV
  - only reword supporting bullet wording to better match the target job
  - never include professional job roles in this section
Do not produce a motivation letter, cover letter, greeting, or signature.
""".strip()
    if extra_instructions:
        prompt = f"{prompt}\n\nAdditional user preferences:\n{extra_instructions.strip()}"
    return prompt


def call_deepseek_json(
    api_key: str,
    model: str,
    prompt: str,
    timeout_seconds: int = 180,
) -> Dict:
    endpoint = "https://api.deepseek.com/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are an expert career writing assistant."},
            {"role": "user", "content": prompt},
        ],
        "stream": False,
    }

    response = requests.post(endpoint, headers=headers, json=payload, timeout=timeout_seconds)
    response.raise_for_status()
    body = response.json()
    choices = body.get("choices") or []
    if not choices:
        raise ValueError("DeepSeek response missing choices.")
    message = choices[0].get("message") or {}
    content = message.get("content") or ""
    return json.loads(strip_json_fences(str(content)))


def generate_docs_for_job(
    deepseek_api_key: Optional[str],
    deepseek_model: str,
    gemini_client: Optional[genai.Client],
    gemini_model: str,
    cv_text: str,
    job: Dict,
    candidate_name: str,
    extra_instructions: str,
    prompt_override: str,
    retries: int,
    retry_sleep: float,
) -> Dict:
    prompt = build_docs_prompt(
        cv_text=cv_text,
        job=job,
        candidate_name=candidate_name,
        extra_instructions=extra_instructions,
        prompt_override=prompt_override,
    )
    last_error = None

    for attempt in range(1, retries + 1):
        try:
            parsed = None
            deepseek_error = None
            gemini_error = None

            if deepseek_api_key:
                try:
                    parsed = call_deepseek_json(
                        api_key=deepseek_api_key,
                        model=deepseek_model,
                        prompt=prompt,
                    )
                except Exception as exc:
                    deepseek_error = exc
                    print(
                        f"DeepSeek failed for job {job.get('job_id')} (attempt {attempt}/{retries}): {exc}. "
                        "Trying Gemini fallback."
                    )

            if parsed is None and gemini_client is not None:
                try:
                    response = gemini_client.models.generate_content(model=gemini_model, contents=prompt)
                    response_text = getattr(response, "text", "") or ""
                    parsed = json.loads(strip_json_fences(response_text))
                except Exception as exc:
                    gemini_error = exc

            if parsed is None:
                if deepseek_error and gemini_error:
                    raise RuntimeError(
                        f"DeepSeek and Gemini fallback both failed. DeepSeek={deepseek_error} | Gemini={gemini_error}"
                    )
                if deepseek_error:
                    raise RuntimeError(f"DeepSeek failed and Gemini fallback is unavailable. DeepSeek={deepseek_error}")
                raise RuntimeError("No Stage 4 AI provider available (DeepSeek/Gemini).")

            professional_summary = str(parsed.get("professional_summary", "")).strip()
            skills_raw = parsed.get("skills", [])
            experiences_raw = parsed.get("professional_experience", [])
            education_raw = parsed.get("education", [])

            if isinstance(skills_raw, list):
                skills = [str(item).strip() for item in skills_raw if str(item).strip()]
            else:
                skills = split_bullets(str(skills_raw))

            experiences = []
            if isinstance(experiences_raw, list):
                for item in experiences_raw:
                    if not isinstance(item, dict):
                        continue
                    role_title = str(item.get("role_title", "")).strip()
                    company = str(item.get("company", "")).strip()
                    period = str(item.get("period", "")).strip()
                    bullets_raw = item.get("bullets", [])
                    if isinstance(bullets_raw, list):
                        bullets = [str(b).strip() for b in bullets_raw if str(b).strip()]
                    else:
                        bullets = split_bullets(str(bullets_raw))
                    if role_title or company or period or bullets:
                        experiences.append(
                            {
                                "role_title": role_title,
                                "company": company,
                                "period": period,
                                "bullets": bullets,
                            }
                        )

            education_entries = []
            if isinstance(education_raw, list):
                for item in education_raw:
                    if not isinstance(item, dict):
                        continue
                    degree_title = str(item.get("degree_title", "")).strip()
                    thesis_title = str(item.get("thesis_title", "")).strip()
                    thesis_bullets_raw = item.get("thesis_bullets", [])
                    if isinstance(thesis_bullets_raw, list):
                        thesis_bullets = [str(b).strip() for b in thesis_bullets_raw if str(b).strip()]
                    else:
                        thesis_bullets = split_bullets(str(thesis_bullets_raw))
                    if degree_title or thesis_title or thesis_bullets:
                        education_entries.append(
                            {
                                "degree_title": degree_title,
                                "thesis_title": thesis_title,
                                "thesis_bullets": thesis_bullets,
                            }
                        )

            if not professional_summary or not skills or not experiences:
                raise ValueError("AI response missing required structured CV fields.")

            tailored_cv_lines = [f"Professional Summary: {professional_summary}", "", "Professional Experience:"]
            for exp in experiences:
                header_parts = [exp.get("role_title", ""), exp.get("company", ""), exp.get("period", "")]
                line = " | ".join([part for part in header_parts if part])
                if line:
                    tailored_cv_lines.append(line)
                for bullet in exp.get("bullets", []):
                    tailored_cv_lines.append(f"- {bullet}")
            tailored_cv_lines.append("")
            tailored_cv_lines.append("Skills:")
            tailored_cv_lines.extend([f"- {skill}" for skill in skills])
            if education_entries:
                tailored_cv_lines.append("")
                tailored_cv_lines.append("Education:")
                for edu in education_entries:
                    if edu.get("degree_title"):
                        tailored_cv_lines.append(str(edu["degree_title"]))
                    if edu.get("thesis_title"):
                        tailored_cv_lines.append(str(edu["thesis_title"]))
                    for bullet in edu.get("thesis_bullets", []):
                        tailored_cv_lines.append(f"- {bullet}")

            return {
                "cv_professional_summary": professional_summary,
                "cv_professional_experience": experiences,
                "cv_education": education_entries,
                "cv_skills": skills,
                "tailored_cv": "\n".join(line for line in tailored_cv_lines if line is not None).strip(),
            }
        except Exception as exc:
            last_error = exc
            if attempt < retries:
                wait_seconds = retry_sleep * attempt
                print(f"Retry {attempt}/{retries - 1} for job {job.get('job_id')} after error: {exc}")
                time.sleep(wait_seconds)

    raise RuntimeError(f"Document generation failed after {retries} attempts: {last_error}")


def split_bullets(text: str) -> List[str]:
    lines = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^[\-\*\u2022]+\s*", "", line)
        lines.append(line)
    return lines


def split_paragraphs(text: str) -> List[str]:
    parts = [item.strip() for item in re.split(r"\n\s*\n", text or "") if item.strip()]
    if parts:
        return parts
    one_block = (text or "").strip()
    return [one_block] if one_block else []


def parse_cv_role_header(role_line: str, fallback_company: str = "") -> Dict:
    cleaned = compact_whitespace(role_line)
    if not cleaned:
        return {"role_title": "", "company": fallback_company, "period": ""}

    parts = [part.strip() for part in cleaned.split("|") if part.strip()]
    role_company_part = parts[0] if parts else cleaned
    period = parts[-1] if len(parts) > 1 else ""

    role_title = ""
    company = fallback_company

    at_match = re.match(r"^(.*?)\s+at\s+(.+)$", role_company_part, flags=re.IGNORECASE)
    if at_match:
        role_title = at_match.group(1).strip(" -")
        company = at_match.group(2).strip()
    else:
        dash_parts = [part.strip() for part in role_company_part.split(" - ", 1)]
        if len(dash_parts) == 2:
            role_title, company = dash_parts[0], dash_parts[1]
        else:
            role_title = role_company_part

    return {
        "role_title": role_title,
        "company": company,
        "period": period,
    }


def normalize_compare_token(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def looks_like_location_value(value: str) -> bool:
    token = normalize_compare_token(value)
    return token in {
        "germany",
        "munich",
        "frankfurt",
        "berlin",
        "hamburg",
        "stuttgart",
        "cologne",
        "dusseldorf",
        "freiburg",
    }


def extract_cv_professional_experiences(cv_text: str) -> List[Dict]:
    lines = [line.rstrip() for line in (cv_text or "").splitlines()]
    if not lines:
        return []

    start_index = None
    for index, line in enumerate(lines):
        if normalize_compare_token(line) == "professional experience":
            start_index = index + 1
            break
    if start_index is None:
        return []

    end_markers = {
        "skills",
        "education",
        "languages",
        "projects",
        "project",
        "strategic technology initiatives",
        "strategic and technology initiatives",
    }
    end_index = len(lines)
    for index in range(start_index, len(lines)):
        marker = normalize_compare_token(lines[index])
        if marker in end_markers:
            end_index = index
            break

    experiences: List[Dict] = []
    cursor = start_index
    while cursor < end_index:
        current = lines[cursor].strip()
        if not current:
            cursor += 1
            continue

        is_header = "|" in current and not re.match(r"^[\-\*\u2022]\s*", current)
        if not is_header:
            cursor += 1
            continue

        header = parse_cv_role_header(current)
        bullets: List[str] = []
        cursor += 1
        while cursor < end_index:
            raw = lines[cursor].strip()
            if not raw:
                if bullets:
                    break
                cursor += 1
                continue
            if "|" in raw and not re.match(r"^[\-\*\u2022]\s*", raw):
                break
            if re.match(r"^[\-\*\u2022]\s*", raw):
                bullets.append(re.sub(r"^[\-\*\u2022]+\s*", "", raw))
            elif bullets:
                bullets[-1] = f"{bullets[-1]} {compact_whitespace(raw)}"
            cursor += 1

        experiences.append(
            {
                "role_title": header["role_title"],
                "company": header["company"],
                "period": header["period"],
                "bullets": bullets,
            }
        )

    return experiences


def match_to_cv_experience(item: Dict, cv_experiences: List[Dict]) -> Optional[Dict]:
    role = str(item.get("role_title") or "").strip()
    company = str(item.get("company") or "").strip()
    period = str(item.get("period") or "").strip()

    if role and " at " in role.lower():
        parsed = parse_cv_role_header(role, fallback_company=company)
        role = parsed.get("role_title", role).strip()
        parsed_company = parsed.get("company", "").strip()
        if parsed_company:
            company = parsed_company
        if not period:
            period = parsed.get("period", "").strip()

    role_key = normalize_compare_token(role)
    company_key = normalize_compare_token(company)
    period_key = normalize_compare_token(period)

    best_score = 0
    best_match = None

    for candidate in cv_experiences:
        candidate_role_key = normalize_compare_token(str(candidate.get("role_title") or ""))
        candidate_company_key = normalize_compare_token(str(candidate.get("company") or ""))
        candidate_period_key = normalize_compare_token(str(candidate.get("period") or ""))
        score = 0

        if role_key and candidate_role_key and (role_key == candidate_role_key or role_key in candidate_role_key or candidate_role_key in role_key):
            score += 5
        if company_key and candidate_company_key and (
            company_key == candidate_company_key or company_key in candidate_company_key or candidate_company_key in company_key
        ):
            score += 4
        if period_key and candidate_period_key and (
            period_key == candidate_period_key or period_key in candidate_period_key or candidate_period_key in period_key
        ):
            score += 3

        if score > best_score:
            best_score = score
            best_match = candidate

    return best_match if best_score >= 5 else None


def dedupe_experiences(experiences: List[Dict]) -> List[Dict]:
    seen = {}
    ordered: List[Dict] = []
    for item in experiences:
        key = (
            normalize_compare_token(str(item.get("role_title") or "")),
            normalize_compare_token(str(item.get("company") or "")),
            normalize_compare_token(str(item.get("period") or "")),
        )
        if key in seen:
            existing = seen[key]
            existing_bullets = [str(b).strip() for b in existing.get("bullets", []) if str(b).strip()]
            existing_keys = {normalize_compare_token(text) for text in existing_bullets}
            for bullet in item.get("bullets", []):
                cleaned = str(bullet).strip()
                if not cleaned:
                    continue
                bullet_key = normalize_compare_token(cleaned)
                if bullet_key and bullet_key not in existing_keys:
                    existing_bullets.append(cleaned)
                    existing_keys.add(bullet_key)
            existing["bullets"] = existing_bullets
            continue

        clone = {
            "role_title": str(item.get("role_title") or "").strip(),
            "company": str(item.get("company") or "").strip(),
            "period": str(item.get("period") or "").strip(),
            "bullets": [str(b).strip() for b in item.get("bullets", []) if str(b).strip()],
        }
        seen[key] = clone
        ordered.append(clone)

    return ordered


def extract_cv_strategic_initiatives(cv_text: str) -> List[Dict]:
    lines = [line.rstrip() for line in (cv_text or "").splitlines()]
    if not lines:
        return []

    start_index = None
    for index, line in enumerate(lines):
        marker = normalize_compare_token(line)
        if marker in {"projects", "project", "strategic technology initiatives", "strategic and technology initiatives"}:
            start_index = index + 1
            break
    if start_index is None:
        return []

    initiatives: List[Dict] = []
    current = None

    for raw in lines[start_index:]:
        line = raw.strip()
        if not line:
            continue
        if normalize_compare_token(line) in {"skills", "education", "professional experience", "languages"}:
            break

        if re.match(r"^[\-\*\u2022]+\s*", line):
            bullet = re.sub(r"^[\-\*\u2022]+\s*", "", line).strip()
            if bullet and current is not None:
                current["bullets"].append(bullet)
            continue

        if current is not None:
            initiatives.append(current)
        current = {"title": compact_whitespace(line), "bullets": []}

    if current is not None:
        initiatives.append(current)

    return [item for item in initiatives if item.get("title")]


def merge_unique_bullets(existing: List[str], additions: List[str]) -> List[str]:
    merged = [str(item).strip() for item in (existing or []) if str(item).strip()]
    seen = {normalize_compare_token(item) for item in merged}
    for bullet in additions or []:
        cleaned = str(bullet).strip()
        if not cleaned:
            continue
        key = normalize_compare_token(cleaned)
        if key and key not in seen:
            merged.append(cleaned)
            seen.add(key)
    return merged


def extract_cv_education(cv_text: str) -> List[Dict]:
    lines = [line.rstrip() for line in (cv_text or "").splitlines()]
    if not lines:
        return []

    start_index = None
    for index, line in enumerate(lines):
        if normalize_compare_token(line) == "education":
            start_index = index + 1
            break
    if start_index is None:
        return []

    end_markers = {
        "projects",
        "project",
        "strategic technology initiatives",
        "strategic and technology initiatives",
        "professional experience",
        "skills",
        "languages",
    }
    end_index = len(lines)
    for index in range(start_index, len(lines)):
        marker = normalize_compare_token(lines[index])
        if marker in end_markers:
            end_index = index
            break

    education_items: List[Dict] = []
    current_item = None
    in_thesis = False

    for cursor in range(start_index, end_index):
        line = lines[cursor].strip()
        if not line:
            continue

        is_degree_line = (
            not line.lower().startswith("master thesis")
            and re.match(
                r"(?i)^(m\.a\.|m\.sc\.|msc\b|b\.sc\.|bsc\b|b\.a\.|mba\b|master\b|bachelor\b|ph\.d\b|phd\b)",
                line,
            )
        )
        if is_degree_line:
            if current_item is not None:
                education_items.append(current_item)
            current_item = {
                "degree_title": compact_whitespace(line),
                "thesis_title": "",
                "thesis_bullets": [],
            }
            in_thesis = False
            continue

        if line.lower().startswith("master thesis"):
            if current_item is None:
                current_item = {"degree_title": "", "thesis_title": "", "thesis_bullets": []}
            current_item["thesis_title"] = compact_whitespace(line)
            in_thesis = True
            continue

        if re.match(r"^(?:[\-\*\u2022]+|--)\s*", line):
            bullet = re.sub(r"^(?:[\-\*\u2022]+|--)\s*", "", line).strip()
            if bullet and current_item is not None:
                current_item["thesis_bullets"] = merge_unique_bullets(current_item.get("thesis_bullets", []), [bullet])
            continue

        if line.lower().startswith("github:"):
            if current_item is not None and in_thesis:
                current_item["thesis_bullets"] = merge_unique_bullets(current_item.get("thesis_bullets", []), [line])
            continue

        if in_thesis and current_item is not None and current_item.get("thesis_bullets"):
            current_item["thesis_bullets"][-1] = compact_whitespace(
                f"{current_item['thesis_bullets'][-1]} {line}"
            )
            continue

        if current_item is not None:
            education_items.append(current_item)
        current_item = {
            "degree_title": compact_whitespace(line),
            "thesis_title": "",
            "thesis_bullets": [],
        }
        in_thesis = False

    if current_item is not None:
        education_items.append(current_item)

    return [item for item in education_items if item.get("degree_title") or item.get("thesis_title")]


def extract_role_from_cv_text(cv_text: str, company_keyword: str):
    lines = [line.rstrip() for line in (cv_text or "").splitlines()]
    keyword = company_keyword.lower()
    for index, line in enumerate(lines):
        if keyword not in line.lower():
            continue

        header = parse_cv_role_header(line, fallback_company=company_keyword)
        bullets = []
        cursor = index + 1
        while cursor < len(lines):
            raw = lines[cursor].strip()
            if not raw:
                if bullets:
                    break
                cursor += 1
                continue

            if re.match(r"^[\-\*\u2022]\s*", raw):
                bullets.append(re.sub(r"^[\-\*\u2022]\s*", "", raw))
                cursor += 1
                continue

            if "|" in raw and not re.match(r"^[\-\*\u2022]\s*", raw):
                break

            if bullets:
                bullets[-1] = f"{bullets[-1]} {compact_whitespace(raw)}"
            cursor += 1

        if bullets:
            return {
                "role_title": header["role_title"],
                "company": header["company"],
                "period": header["period"],
                "bullets": bullets,
            }
    return None


def ensure_structured_cv_fields(record: Dict, candidate_name: str, cv_text: str) -> None:
    title = str(record.get("title") or "").strip()
    company = str(record.get("company") or "").strip()

    if not record.get("cv_professional_summary"):
        record["cv_professional_summary"] = (
            f"{candidate_name} applying to {title} at {company}, with focus on business transformation, "
            "AI implementation, and measurable delivery impact."
        ).strip()

    skills = record.get("cv_skills", [])
    if isinstance(skills, list):
        normalized_skills = [str(item).strip() for item in skills if str(item).strip()]
    else:
        normalized_skills = split_bullets(str(skills))
    record["cv_skills"] = normalized_skills

    cv_experiences = extract_cv_professional_experiences(cv_text)
    experience_bullets_by_key: Dict[tuple, List[str]] = {}
    experiences = record.get("cv_professional_experience", [])
    if isinstance(experiences, list):
        for item in experiences:
            if not isinstance(item, dict):
                continue
            matched_cv_experience = match_to_cv_experience(item, cv_experiences)
            if not matched_cv_experience:
                continue
            bullets_raw = item.get("bullets", [])
            if isinstance(bullets_raw, list):
                bullets = [str(b).strip() for b in bullets_raw if str(b).strip()]
            else:
                bullets = split_bullets(str(bullets_raw))
            key = (
                normalize_compare_token(str(matched_cv_experience.get("role_title", ""))),
                normalize_compare_token(str(matched_cv_experience.get("company", ""))),
                normalize_compare_token(str(matched_cv_experience.get("period", ""))),
            )
            if key not in experience_bullets_by_key:
                experience_bullets_by_key[key] = []
            experience_bullets_by_key[key] = merge_unique_bullets(experience_bullets_by_key[key], bullets)

    normalized_experiences = []
    if cv_experiences:
        for base_item in cv_experiences:
            key = (
                normalize_compare_token(str(base_item.get("role_title", ""))),
                normalize_compare_token(str(base_item.get("company", ""))),
                normalize_compare_token(str(base_item.get("period", ""))),
            )
            base_bullets = [str(b).strip() for b in base_item.get("bullets", []) if str(b).strip()]
            selected_bullets = experience_bullets_by_key.get(key) or base_bullets
            normalized_experiences.append(
                {
                    "role_title": str(base_item.get("role_title", "")).strip(),
                    "company": str(base_item.get("company", "")).strip(),
                    "period": str(base_item.get("period", "")).strip(),
                    "bullets": selected_bullets,
                }
            )
    else:
        normalized_experiences = dedupe_experiences(
            [
                {
                    "role_title": str(item.get("role_title", "")).strip(),
                    "company": str(item.get("company", "")).strip(),
                    "period": str(item.get("period", "")).strip(),
                    "bullets": [str(b).strip() for b in item.get("bullets", []) if str(b).strip()],
                }
                for item in experiences
                if isinstance(item, dict)
            ]
        )

    record["cv_professional_experience"] = normalized_experiences

    baseline_education = extract_cv_education(cv_text)
    generated_education = record.get("cv_education", [])
    normalized_generated_education = []
    if isinstance(generated_education, list):
        for item in generated_education:
            if not isinstance(item, dict):
                continue
            degree_title = compact_whitespace(str(item.get("degree_title", "")).strip())
            thesis_title = compact_whitespace(str(item.get("thesis_title", "")).strip())
            thesis_raw = item.get("thesis_bullets", [])
            if isinstance(thesis_raw, list):
                thesis_bullets = [str(b).strip() for b in thesis_raw if str(b).strip()]
            else:
                thesis_bullets = split_bullets(str(thesis_raw))
            if degree_title or thesis_title or thesis_bullets:
                normalized_generated_education.append(
                    {
                        "degree_title": degree_title,
                        "thesis_title": thesis_title,
                        "thesis_bullets": thesis_bullets,
                    }
                )

    if baseline_education:
        final_education = []
        for base_item in baseline_education:
            base_degree = compact_whitespace(str(base_item.get("degree_title", "")).strip())
            base_thesis_title = compact_whitespace(str(base_item.get("thesis_title", "")).strip())
            base_thesis_bullets = [str(b).strip() for b in base_item.get("thesis_bullets", []) if str(b).strip()]
            base_key = normalize_compare_token(base_degree)

            matched_generated = None
            for gen_item in normalized_generated_education:
                gen_key = normalize_compare_token(str(gen_item.get("degree_title", "")))
                if not gen_key:
                    continue
                if gen_key == base_key or gen_key in base_key or base_key in gen_key:
                    matched_generated = gen_item
                    break

            thesis_bullets = base_thesis_bullets
            if matched_generated and matched_generated.get("thesis_bullets"):
                thesis_bullets = merge_unique_bullets([], matched_generated.get("thesis_bullets", []))

            final_education.append(
                {
                    "degree_title": base_degree,
                    "thesis_title": base_thesis_title,
                    "thesis_bullets": thesis_bullets,
                }
            )
        record["cv_education"] = final_education
    else:
        record["cv_education"] = normalized_generated_education

    baseline_initiatives = extract_cv_strategic_initiatives(cv_text)
    generated_initiatives = record.get("cv_strategic_initiatives", [])
    normalized_generated_initiatives: List[Dict] = []
    if isinstance(generated_initiatives, list):
        for item in generated_initiatives:
            if isinstance(item, dict):
                item_title = compact_whitespace(str(item.get("title", "")).strip())
                bullets_raw = item.get("bullets", [])
                if isinstance(bullets_raw, list):
                    item_bullets = [str(b).strip() for b in bullets_raw if str(b).strip()]
                else:
                    item_bullets = split_bullets(str(bullets_raw))
                if item_title:
                    normalized_generated_initiatives.append({"title": item_title, "bullets": item_bullets})
            elif isinstance(item, str):
                item_title = compact_whitespace(item)
                if item_title:
                    normalized_generated_initiatives.append({"title": item_title, "bullets": []})

    if baseline_initiatives:
        initiatives_by_baseline_title: Dict[str, List[str]] = {}
        for gen_item in normalized_generated_initiatives:
            gen_title_key = normalize_compare_token(str(gen_item.get("title", "")))
            if not gen_title_key:
                continue
            for base_item in baseline_initiatives:
                base_title = compact_whitespace(str(base_item.get("title", "")))
                base_key = normalize_compare_token(base_title)
                if not base_key:
                    continue
                if gen_title_key == base_key or gen_title_key in base_key or base_key in gen_title_key:
                    existing = initiatives_by_baseline_title.get(base_key, [])
                    initiatives_by_baseline_title[base_key] = merge_unique_bullets(existing, gen_item.get("bullets", []))
                    break

        final_initiatives = []
        for base_item in baseline_initiatives:
            base_title = compact_whitespace(str(base_item.get("title", "")))
            base_key = normalize_compare_token(base_title)
            base_bullets = [str(b).strip() for b in base_item.get("bullets", []) if str(b).strip()]
            selected_bullets = initiatives_by_baseline_title.get(base_key) or base_bullets
            final_initiatives.append({"title": base_title, "bullets": selected_bullets})
        record["cv_strategic_initiatives"] = final_initiatives
    else:
        deduped_initiatives = []
        seen_initiatives = set()
        for item in normalized_generated_initiatives:
            title_key = normalize_compare_token(str(item.get("title", "")))
            if not title_key or title_key in seen_initiatives:
                continue
            seen_initiatives.add(title_key)
            deduped_initiatives.append(
                {
                    "title": compact_whitespace(str(item.get("title", ""))),
                    "bullets": [str(b).strip() for b in item.get("bullets", []) if str(b).strip()],
                }
            )
        record["cv_strategic_initiatives"] = deduped_initiatives


def resolve_profile_image_path(raw_path: str):
    candidate = normalize_windows_env_path(raw_path)
    if not candidate:
        return None
    image_path = Path(candidate)
    if image_path.suffix.lower() != ".png":
        return None
    return image_path if image_path.exists() and image_path.is_file() else None


def resolve_assets_profile_png(docs_dir: Path):
    assets_dir = docs_dir / "_assets"
    preferred = assets_dir / "_profile_from_cv.png"
    if preferred.exists() and preferred.is_file():
        return preferred
    png_files = sorted(assets_dir.glob("*.png")) if assets_dir.exists() else []
    return png_files[0] if png_files else None


def find_cv_docx_source_path():
    candidates = []
    config = load_job_seeker_config()
    config_cv_path = normalize_windows_env_path(cfg_str(config, ("candidate", "cv_path"), ""))
    config_cv_docx_path = normalize_windows_env_path(cfg_str(config, ("candidate", "cv_docx_path"), ""))
    if config_cv_path:
        candidates.append(Path(config_cv_path))
    if config_cv_docx_path:
        candidates.append(Path(config_cv_docx_path))
    env_cv_path = normalize_windows_env_path(os.getenv("MY_CV_PATH", ""))
    if env_cv_path:
        candidates.append(Path(env_cv_path))
    candidates.append(Path("Ahmed Kaddah CV.docx"))
    candidates.append(Path(r"C:\Users\ahmed\OneDrive\Personal\CV\Ahmed Kaddah CV.docx"))

    for path in candidates:
        if path.exists() and path.is_file() and path.suffix.lower() == ".docx":
            return path
    return None


def extract_profile_image_from_cv_docx(cv_docx_path: Path, docs_dir: Path):
    try:
        docs_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(cv_docx_path, "r") as archive:
            media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
            if not media_files:
                return None
            media_files.sort()
            first_media = media_files[0]
            extension = Path(first_media).suffix or ".png"
            output_path = docs_dir / f"_profile_from_cv{extension}"
            output_path.write_bytes(archive.read(first_media))
            return output_path if output_path.exists() else None
    except Exception:
        return None


def create_cv_document(
    record: Dict,
    docs_dir: Path,
    run_date: str,
    candidate_name: str,
    candidate_email: str,
    cv_font_name: str,
    languages: List[str],
    profile_image_path,
) -> str:
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("python-docx is required to create Word files. Install: pip install python-docx") from exc

    job_id = str(record.get("job_id", "unknown"))
    title = str(record.get("title", "Untitled"))
    company = str(record.get("company", "Unknown Company"))
    header_location = format_header_location(record)
    safe_stem = sanitize_filename(f"{candidate_name}_{title}_{company}_{job_id}_CV", max_length=140)

    target_dir = docs_dir / run_date
    target_dir.mkdir(parents=True, exist_ok=True)

    cv_path = target_dir / f"{safe_stem}.docx"

    doc = Document()
    for style_name in ("Normal", "Heading 1", "Heading 2", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = cv_font_name
        style.font.size = Pt(10.5)
        try:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), cv_font_name)
        except Exception:
            pass

    section = doc.sections[0]
    margin = Inches(0.1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    name_paragraph = doc.add_paragraph()
    name_paragraph.paragraph_format.space_before = Pt(0)
    name_paragraph.paragraph_format.space_after = Pt(2)
    language_values = [str(line).strip() for line in (languages or DEFAULT_LANGUAGES) if str(line).strip()]
    language_line = ", ".join(language_values)
    header_parts = [candidate_name, header_location, candidate_email]
    if language_line:
        header_parts.append(language_line)
    name_run = name_paragraph.add_run(" | ".join([part for part in header_parts if part]))
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.name = cv_font_name

    def float_picture_right(run, inline_shape, top_offset_inches: float = 0.45):
        inline = inline_shape._inline
        drawing = run._r.xpath("./w:drawing")[0]

        anchor = OxmlElement("wp:anchor")
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "251658240")
        anchor.set("behindDoc", "0")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")
        anchor.set("distT", "0")
        anchor.set("distB", "0")
        anchor.set("distL", "91440")
        anchor.set("distR", "91440")

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        anchor.append(simple_pos)

        position_h = OxmlElement("wp:positionH")
        position_h.set("relativeFrom", "margin")
        align = OxmlElement("wp:align")
        align.text = "right"
        position_h.append(align)
        anchor.append(position_h)

        position_v = OxmlElement("wp:positionV")
        position_v.set("relativeFrom", "margin")
        pos_offset = OxmlElement("wp:posOffset")
        pos_offset.text = str(int(top_offset_inches * 914400))
        position_v.append(pos_offset)
        anchor.append(position_v)

        extent = inline.xpath("./wp:extent")[0]
        anchor.append(deepcopy(extent))

        effect = inline.xpath("./wp:effectExtent")
        if effect:
            anchor.append(deepcopy(effect[0]))
        else:
            effect_extent = OxmlElement("wp:effectExtent")
            effect_extent.set("l", "0")
            effect_extent.set("t", "0")
            effect_extent.set("r", "0")
            effect_extent.set("b", "0")
            anchor.append(effect_extent)

        wrap_square = OxmlElement("wp:wrapSquare")
        wrap_square.set("wrapText", "bothSides")
        anchor.append(wrap_square)

        doc_pr = inline.xpath("./wp:docPr")[0]
        anchor.append(deepcopy(doc_pr))

        frame_pr = inline.xpath("./wp:cNvGraphicFramePr")
        if frame_pr:
            anchor.append(deepcopy(frame_pr[0]))

        graphic = inline.xpath("./a:graphic")[0]
        anchor.append(deepcopy(graphic))

        drawing.remove(inline)
        drawing.append(anchor)

    if profile_image_path:
        try:
            image_run = name_paragraph.add_run()
            inline_shape = image_run.add_picture(str(profile_image_path), width=Inches(1.5))
            float_picture_right(image_run, inline_shape, top_offset_inches=0.45)
        except Exception:
            pass

    def add_section_separator() -> None:
        # Use a subtle horizontal rule instead of long dash text.
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.space_after = Pt(0)
        sep.paragraph_format.line_spacing = Pt(1)
        p_pr = sep._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    heading_summary = doc.add_paragraph()
    run_summary = heading_summary.add_run("PROFESSIONAL SUMMARY")
    run_summary.bold = True
    run_summary.font.color.rgb = None
    doc.add_paragraph(str(record.get("cv_professional_summary") or "").strip())
    add_section_separator()

    heading_exp = doc.add_paragraph()
    run_exp = heading_exp.add_run("PROFESSIONAL EXPERIENCE")
    run_exp.bold = True
    run_exp.font.color.rgb = None

    experiences = record.get("cv_professional_experience") or []
    for item in experiences:
        if not isinstance(item, dict):
            continue
        role_title = str(item.get("role_title") or "").strip()
        exp_company = str(item.get("company") or "").strip()
        period = str(item.get("period") or "").strip()
        headline_parts = [part for part in [role_title, exp_company, period] if part]
        if headline_parts:
            exp_header = doc.add_paragraph(" | ".join(headline_parts))
            exp_header.runs[0].bold = True
        for bullet in item.get("bullets", []):
            if str(bullet).strip():
                doc.add_paragraph(str(bullet).strip(), style="List Bullet")
    add_section_separator()

    heading_initiatives = doc.add_paragraph()
    run_initiatives = heading_initiatives.add_run("PROJECTS")
    run_initiatives.bold = True
    initiatives = record.get("cv_strategic_initiatives") or []
    for item in initiatives:
        if not isinstance(item, dict):
            continue
        initiative_title = str(item.get("title") or "").strip()
        if initiative_title:
            ini_header = doc.add_paragraph(initiative_title)
            ini_header.runs[0].bold = True
        for bullet in item.get("bullets", []):
            if str(bullet).strip():
                doc.add_paragraph(str(bullet).strip(), style="List Bullet")
    add_section_separator()

    heading_skills = doc.add_paragraph()
    run_skills = heading_skills.add_run("SKILLS")
    run_skills.bold = True
    skills = record.get("cv_skills") or []
    if skills:
        doc.add_paragraph(", ".join([str(skill).strip() for skill in skills if str(skill).strip()]))
    add_section_separator()

    heading_education = doc.add_paragraph()
    run_education = heading_education.add_run("EDUCATION")
    run_education.bold = True
    education_items = record.get("cv_education") or []
    for item in education_items:
        if not isinstance(item, dict):
            continue
        degree_title = str(item.get("degree_title") or "").strip()
        thesis_title = str(item.get("thesis_title") or "").strip()
        if degree_title:
            degree_paragraph = doc.add_paragraph(degree_title)
            degree_paragraph.runs[0].bold = True
        if thesis_title:
            doc.add_paragraph(thesis_title)
        for bullet in item.get("thesis_bullets", []):
            if str(bullet).strip():
                doc.add_paragraph(str(bullet).strip(), style="List Bullet")
    add_section_separator()

    doc.save(cv_path)
    return str(cv_path.resolve())


def convert_docx_to_pdf(docx_path: str) -> str:
    source_path = Path(docx_path)
    target_path = source_path.with_suffix(".pdf")

    try:
        from docx2pdf import convert as docx2pdf_convert

        docx2pdf_convert(str(source_path), str(target_path))
        if target_path.exists():
            return str(target_path.resolve())
    except Exception:
        pass

    try:
        src = str(source_path.resolve()).replace("'", "''")
        dst = str(target_path.resolve()).replace("'", "''")
        ps_script = (
            "$word = New-Object -ComObject Word.Application; "
            "$word.Visible = $false; "
            f"$doc = $word.Documents.Open('{src}'); "
            f"$doc.SaveAs('{dst}', 17); "
            "$doc.Close(); "
            "$word.Quit();"
        )
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps_script],
            check=True,
            capture_output=True,
            text=True,
        )
        if target_path.exists():
            return str(target_path.resolve())
    except Exception:
        pass

    office_cmd = shutil.which("soffice") or shutil.which("libreoffice")
    if office_cmd:
        try:
            subprocess.run(
                [office_cmd, "--headless", "--convert-to", "pdf", "--outdir", str(source_path.parent), str(source_path)],
                check=True,
                capture_output=True,
                text=True,
            )
            if target_path.exists():
                return str(target_path.resolve())
        except Exception:
            pass

    raise RuntimeError(
        "Unable to convert DOCX to PDF. Install docx2pdf + Microsoft Word, "
        "or install LibreOffice and add soffice to PATH."
    )


def to_cell_value(value):
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def derive_columns(records: List[Dict]) -> List[str]:
    preferred_columns = [
        "run_date",
        "run_timestamp",
        "job_id",
        "title",
        "company",
        "location_raw",
        "keyword",
        "posted_time_text",
        "posted_age_hours",
        "applicant_count",
        "priority_rank",
        "priority_rule",
        "easy_apply_status",
        "apply_link",
        "apply_link_source",
        "linkedin_link",
        "link",
        "enrich_status_code",
        "enrich_error",
        "full_description",
        "cv_professional_summary",
        "cv_professional_experience",
        "cv_strategic_initiatives",
        "cv_skills",
        "cv_education",
        "tailored_cv",
        "cv_docx",
        "cv_pdf",
        "tailored_cv_docx",
        "pdf_generation_error",
        "doc_generation_error",
    ]

    seen = set()
    for record in records:
        seen.update(record.keys())

    columns = []
    for column in preferred_columns:
        if column in seen:
            columns.append(column)
            seen.discard(column)

    columns.extend(sorted(seen))
    return columns


def sheet_is_empty(worksheet) -> bool:
    return worksheet.max_row == 1 and worksheet.max_column == 1 and worksheet["A1"].value is None


def add_sheet_with_unique_name(workbook, base_name: str):
    clean_base = (base_name or "jobs").strip() or "jobs"
    clean_base = clean_base[:28]
    candidate = clean_base
    counter = 2
    while candidate in workbook.sheetnames:
        suffix = f"_{counter}"
        candidate = f"{clean_base[:31-len(suffix)]}{suffix}"
        counter += 1
    return workbook.create_sheet(title=candidate)


def apply_hyperlink(cell, raw_value: str) -> None:
    value = (raw_value or "").strip()
    if not value:
        return

    if value.startswith("http://") or value.startswith("https://"):
        cell.hyperlink = value
        cell.style = "Hyperlink"
        return

    path = Path(value)
    if path.exists():
        cell.hyperlink = path.resolve().as_uri()
        cell.style = "Hyperlink"


def style_worksheet(worksheet, headers: List[str]) -> None:
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter

    for header_cell in worksheet[1]:
        header_cell.font = Font(bold=True)
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions

    width_defaults = {
        "run_date": 12,
        "run_timestamp": 22,
        "job_id": 14,
        "title": 42,
        "company": 32,
        "location_raw": 24,
        "keyword": 22,
        "posted_time_text": 20,
        "posted_age_hours": 14,
        "applicant_count": 14,
        "priority_rank": 12,
        "priority_rule": 34,
        "easy_apply_status": 16,
        "apply_link": 45,
        "apply_link_source": 18,
        "linkedin_link": 45,
        "link": 45,
        "enrich_status_code": 16,
        "enrich_error": 30,
        "full_description": 80,
        "cv_professional_summary": 80,
        "cv_professional_experience": 90,
        "cv_strategic_initiatives": 90,
        "cv_skills": 55,
        "cv_education": 80,
        "tailored_cv": 80,
        "cv_docx": 45,
        "cv_pdf": 45,
        "tailored_cv_docx": 45,
        "pdf_generation_error": 35,
        "doc_generation_error": 30,
    }

    for index, header in enumerate(headers, start=1):
        column_letter = get_column_letter(index)
        worksheet.column_dimensions[column_letter].width = width_defaults.get(header, 24)

    wrap_columns = {
        "full_description",
        "cv_professional_summary",
        "cv_professional_experience",
        "cv_strategic_initiatives",
        "cv_skills",
        "cv_education",
        "tailored_cv",
        "pdf_generation_error",
        "enrich_error",
        "doc_generation_error",
    }
    header_to_index = {header: idx for idx, header in enumerate(headers, start=1)}
    for column_name in wrap_columns:
        if column_name not in header_to_index:
            continue
        col_index = header_to_index[column_name]
        column_letter = get_column_letter(col_index)
        for cell in worksheet[column_letter]:
            cell.alignment = Alignment(wrap_text=True, vertical="top")


def write_records_to_worksheet(worksheet, records: List[Dict], headers: List[str], append_only: bool) -> None:
    hyperlink_columns = {
        "apply_link",
        "linkedin_link",
        "link",
        "cv_docx",
        "cv_pdf",
        "tailored_cv_docx",
    }
    header_to_index = {header: idx for idx, header in enumerate(headers, start=1)}

    if not append_only:
        worksheet.append(headers)

    for record in records:
        row_values = [to_cell_value(record.get(header)) for header in headers]
        worksheet.append(row_values)
        row_index = worksheet.max_row
        for column_name in hyperlink_columns:
            col_index = header_to_index.get(column_name)
            if not col_index:
                continue
            cell = worksheet.cell(row=row_index, column=col_index)
            apply_hyperlink(cell, cell.value or "")


def save_to_excel(records: List[Dict], output_path: Path, excel_mode: str, sheet_name: str, run_date: str) -> None:
    try:
        from openpyxl import Workbook, load_workbook
    except Exception as exc:
        raise RuntimeError("openpyxl is required for Excel export. Install with: pip install openpyxl") from exc

    if output_path.exists():
        workbook = load_workbook(output_path)
    else:
        workbook = Workbook()

    headers = derive_columns(records)

    if excel_mode == "append-rows":
        target_name = (sheet_name or "jobs").strip() or "jobs"
        if target_name in workbook.sheetnames:
            worksheet = workbook[target_name]
        else:
            worksheet = workbook.active if sheet_is_empty(workbook.active) else workbook.create_sheet(title=target_name)
            worksheet.title = target_name

        if sheet_is_empty(worksheet):
            write_records_to_worksheet(worksheet, records, headers, append_only=False)
        else:
            existing_headers = [cell.value for cell in worksheet[1] if cell.value]
            merged_headers = list(existing_headers)
            for header in headers:
                if header not in merged_headers:
                    merged_headers.append(header)
            if merged_headers != existing_headers:
                for index, header in enumerate(merged_headers, start=1):
                    worksheet.cell(row=1, column=index, value=header)
                headers = merged_headers
            else:
                headers = existing_headers
            write_records_to_worksheet(worksheet, records, headers, append_only=True)
    else:
        target_sheet_name = sheet_name or run_date
        if sheet_is_empty(workbook.active) and len(workbook.sheetnames) == 1:
            worksheet = workbook.active
            worksheet.title = (target_sheet_name[:31] or "jobs")
        else:
            worksheet = add_sheet_with_unique_name(workbook, target_sheet_name)
        write_records_to_worksheet(worksheet, records, headers, append_only=False)

    for ws in workbook.worksheets:
        current_headers = [cell.value for cell in ws[1] if cell.value]
        if current_headers:
            style_worksheet(ws, current_headers)

    workbook.save(output_path)


def main() -> int:
    load_dotenv()
    config = load_job_seeker_config()

    default_input_json = cfg_str(config, ("runtime", "stage4", "input_json"), "stage3_filtered_ai.json")
    default_output_json = cfg_str(config, ("outputs", "stage4_json"), "stage4_documents.json")
    default_output_xlsx = cfg_str(config, ("outputs", "stage4_xlsx"), "final_jobs_with_docs.xlsx")
    default_docs_dir = cfg_str(config, ("outputs", "docs_dir"), "generated_docs")
    default_stage4_checkpoint = cfg_str(config, ("runtime", "stage4", "checkpoint_json"), "stage4_checkpoint.json")
    default_deepseek_model = cfg_str(
        config,
        ("ai", "models", "stage4_docs_deepseek"),
        os.getenv("DEEPSEEK_STAGE4_MODEL", "deepseek-chat"),
    )
    default_gemini_fallback_model = (
        cfg_str(config, ("ai", "models", "stage4_docs_fallback_gemini"), "")
        or os.getenv("GEMINI_DOCS_MODEL", "gemini-2.5-flash")
    )
    default_candidate_name = cfg_str(config, ("candidate", "name"), "") or os.getenv(
        "CANDIDATE_NAME",
        DEFAULT_CANDIDATE_NAME,
    )
    default_candidate_email = cfg_str(config, ("candidate", "email"), "") or os.getenv(
        "CANDIDATE_EMAIL",
        DEFAULT_CANDIDATE_EMAIL,
    )
    default_profile_image = cfg_str(config, ("candidate", "profile_image_path"), "") or os.getenv(
        "CV_PROFILE_IMAGE_PATH",
        "",
    )
    default_cv_font = cfg_str(config, ("candidate", "cv_font"), DEFAULT_CV_FONT) or DEFAULT_CV_FONT
    default_languages = [str(item) for item in cfg_list(config, ("candidate", "languages"), DEFAULT_LANGUAGES)]
    default_stage4_extra_prompt = cfg_str(config, ("ai", "prompts", "stage4_extra_instructions"), "")
    default_stage4_prompt_override = cfg_str(config, ("ai", "prompts", "stage4_prompt_override"), "")
    default_stage4_sleep_seconds = cfg_float(
        config,
        ("runtime", "stage4", "sleep_seconds"),
        float(os.getenv("STAGE4_SLEEP_SECONDS", "4")),
    )
    default_stage4_retries = cfg_int(
        config,
        ("runtime", "stage4", "retries"),
        int(os.getenv("STAGE4_RETRIES", "3")),
    )
    default_stage4_retry_sleep = cfg_float(
        config,
        ("runtime", "stage4", "retry_sleep_seconds"),
        float(os.getenv("STAGE4_RETRY_SLEEP_SECONDS", "3")),
    )
    default_stage4_max_jobs = cfg_int(
        config,
        ("runtime", "stage4", "max_jobs"),
        int(os.getenv("STAGE4_MAX_JOBS", "0")),
    )
    default_stage4_excel_mode = cfg_str(
        config,
        ("runtime", "stage4", "excel_mode"),
        os.getenv("STAGE4_EXCEL_MODE", "new-sheet"),
    )
    if default_stage4_excel_mode not in ("new-sheet", "append-rows"):
        default_stage4_excel_mode = "new-sheet"
    default_stage4_sheet_name = cfg_str(
        config,
        ("runtime", "stage4", "sheet_name"),
        os.getenv("STAGE4_SHEET_NAME", ""),
    )
    default_stage4_run_date = cfg_str(
        config,
        ("runtime", "stage4", "run_date"),
        os.getenv("STAGE4_RUN_DATE", ""),
    )
    default_stage4_force_regenerate = cfg_bool(
        config,
        ("runtime", "stage4", "force_regenerate"),
        os.getenv("STAGE4_FORCE_REGENERATE", "false").lower() in ("1", "true", "yes"),
    )

    parser = argparse.ArgumentParser(
        description=(
            "Stage 4: generate a structured CV and export both .docx and .pdf per job, "
            "plus JSON/XLSX."
        )
    )
    parser.add_argument("--input", default=default_input_json, help="Input JSON from Stage 3.")
    parser.add_argument("--output-json", default=default_output_json, help="Output JSON with generated documents.")
    parser.add_argument("--output-xlsx", default=default_output_xlsx, help="Output Excel file.")
    parser.add_argument("--checkpoint", default=default_stage4_checkpoint, help="Checkpoint for resumable generation.")
    parser.add_argument("--docs-dir", default=default_docs_dir, help="Directory where .docx files are stored.")
    parser.add_argument(
        "--model",
        default=default_deepseek_model,
        help="Primary DeepSeek model for Stage 4 (e.g., deepseek-reasoner, deepseek-chat).",
    )
    parser.add_argument(
        "--fallback-model",
        default=default_gemini_fallback_model,
        help="Gemini fallback model used only if DeepSeek fails.",
    )
    parser.add_argument(
        "--candidate-name",
        default=default_candidate_name,
        help="Candidate name used in document title, filename, and signature.",
    )
    parser.add_argument(
        "--candidate-email",
        default=default_candidate_email,
        help="Candidate email shown in CV header.",
    )
    parser.add_argument(
        "--profile-image",
        default=default_profile_image,
        help="Optional path to profile image shown on the top-right of the CV header.",
    )
    parser.add_argument(
        "--cv-font",
        default=default_cv_font,
        choices=["Calibri", "Arial"],
        help="CV font family.",
    )
    parser.add_argument(
        "--languages",
        nargs="*",
        default=default_languages,
        help="Language lines printed in LANGUAGES section.",
    )
    parser.add_argument(
        "--stage4-extra-prompt",
        default=default_stage4_extra_prompt,
        help="Extra instructions appended to Stage 4 generation prompt.",
    )
    parser.add_argument(
        "--stage4-prompt-override",
        default=default_stage4_prompt_override,
        help=(
            "Optional full prompt override. Supports placeholders: {{CV_TEXT}}, {{JOB_ID}}, "
            "{{JOB_TITLE}}, {{JOB_COMPANY}}, {{JOB_CITY}}, {{JOB_DESCRIPTION}}, {{CANDIDATE_NAME}}."
        ),
    )
    parser.add_argument("--sleep-seconds", type=float, default=default_stage4_sleep_seconds)
    parser.add_argument("--retries", type=int, default=default_stage4_retries)
    parser.add_argument("--retry-sleep", type=float, default=default_stage4_retry_sleep)
    parser.add_argument(
        "--max-jobs",
        type=int,
        default=default_stage4_max_jobs,
        help="0 means all jobs. Use a positive number to cap AI document generation for quota control.",
    )
    parser.add_argument(
        "--excel-mode",
        choices=["new-sheet", "append-rows"],
        default=default_stage4_excel_mode,
        help="new-sheet: create a dated sheet each run. append-rows: append into one sheet.",
    )
    parser.add_argument(
        "--sheet-name",
        default=default_stage4_sheet_name,
        help="Sheet name. For new-sheet mode, empty means run date; for append-rows, default is jobs.",
    )
    parser.add_argument(
        "--run-date",
        default=default_stage4_run_date,
        help="Override run date (YYYY-MM-DD). Empty uses today.",
    )
    parser.add_argument(
        "--force-regenerate",
        action=argparse.BooleanOptionalAction,
        default=default_stage4_force_regenerate,
        help="Ignore existing stage4 checkpoint and regenerate docs for all selected jobs.",
    )
    args = parser.parse_args()

    deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not deepseek_api_key and not gemini_api_key:
        print("ERROR: Both DEEPSEEK_API_KEY and GEMINI_API_KEY are missing in environment/.env")
        return 1

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}")
        return 1

    jobs = load_json_file(input_path)
    if not isinstance(jobs, list):
        print("ERROR: Input JSON must be a list of jobs.")
        return 1

    if args.max_jobs > 0:
        jobs = jobs[: args.max_jobs]

    run_dt = datetime.now()
    run_date = args.run_date.strip() or run_dt.strftime("%Y-%m-%d")
    run_timestamp = run_dt.isoformat(timespec="seconds")
    docs_dir = Path(args.docs_dir)
    candidate_name = (args.candidate_name or DEFAULT_CANDIDATE_NAME).strip() or DEFAULT_CANDIDATE_NAME
    candidate_email = (args.candidate_email or DEFAULT_CANDIDATE_EMAIL).strip() or DEFAULT_CANDIDATE_EMAIL
    cv_font_name = (args.cv_font or DEFAULT_CV_FONT).strip() or DEFAULT_CV_FONT
    languages = [str(item).strip() for item in (args.languages or DEFAULT_LANGUAGES) if str(item).strip()]
    if not languages:
        languages = list(DEFAULT_LANGUAGES)
    normalized_profile_input = normalize_windows_env_path(args.profile_image)
    profile_image_path = resolve_profile_image_path(normalized_profile_input)
    if normalized_profile_input and not profile_image_path:
        print(f"WARNING: profile image must be an existing .png file, got: {normalized_profile_input}")
    if not profile_image_path:
        profile_image_path = resolve_assets_profile_png(docs_dir)
        if profile_image_path:
            print(f"INFO: using profile photo from assets PNG: {profile_image_path}")
        else:
            print("WARNING: no PNG profile image found in docs _assets folder.")

    cv_text = load_cv_text()
    gemini_client = genai.Client(api_key=gemini_api_key) if gemini_api_key else None

    checkpoint_path = Path(args.checkpoint)
    checkpoint = {"generated_records": []}
    if checkpoint_path.exists() and not args.force_regenerate:
        loaded_checkpoint = load_json_file(checkpoint_path)
        if isinstance(loaded_checkpoint, dict):
            checkpoint.update(loaded_checkpoint)

    generated_by_id = {}
    for record in checkpoint.get("generated_records", []):
        job_id = str(record.get("job_id"))
        if job_id:
            generated_by_id[job_id] = record

    total_jobs = len(jobs)

    checkpoint_changed = False

    for index, job in enumerate(jobs, start=1):
        job_id = str(job.get("job_id"))
        if job_id in generated_by_id and not args.force_regenerate:
            existing_record = generated_by_id[job_id]
            for passthrough_key in [
                "posted_time_text",
                "posted_age_hours",
                "posted_datetime_estimated_utc",
                "applicant_count",
                "priority_rank",
                "priority_tier",
                "priority_bucket",
                "priority_rule",
            ]:
                if passthrough_key in job and existing_record.get(passthrough_key) != job.get(passthrough_key):
                    existing_record[passthrough_key] = job.get(passthrough_key)
                    checkpoint_changed = True
            if not existing_record.get("run_date"):
                existing_record["run_date"] = run_date
                checkpoint_changed = True
            if not existing_record.get("run_timestamp"):
                existing_record["run_timestamp"] = run_timestamp
                checkpoint_changed = True
            if not existing_record.get("linkedin_link"):
                existing_record["linkedin_link"] = existing_record.get("link", "")
                checkpoint_changed = True
            if not existing_record.get("apply_link"):
                existing_record["apply_link"] = existing_record.get("linkedin_link", "")
                checkpoint_changed = True

            ensure_structured_cv_fields(existing_record, candidate_name=candidate_name, cv_text=cv_text)

            missing_cv_doc = not existing_record.get("cv_docx")
            has_text_content = bool(existing_record.get("cv_professional_summary")) and bool(
                existing_record.get("cv_professional_experience")
            )
            if missing_cv_doc and has_text_content:
                try:
                    cv_doc_path = create_cv_document(
                        existing_record,
                        docs_dir=docs_dir,
                        run_date=existing_record.get("run_date", run_date),
                        candidate_name=candidate_name,
                        candidate_email=candidate_email,
                        cv_font_name=cv_font_name,
                        languages=languages,
                        profile_image_path=profile_image_path,
                    )
                    existing_record["cv_docx"] = cv_doc_path
                    existing_record["tailored_cv_docx"] = cv_doc_path
                    existing_record["doc_generation_error"] = None
                    checkpoint_changed = True
                except Exception as exc:
                    existing_record["doc_generation_error"] = str(exc)
                    checkpoint_changed = True
            elif missing_cv_doc and existing_record.get("tailored_cv_docx"):
                existing_record["cv_docx"] = existing_record.get("tailored_cv_docx")
                checkpoint_changed = True

            missing_pdf = not existing_record.get("cv_pdf")
            if not missing_pdf and existing_record.get("cv_pdf") and not Path(existing_record["cv_pdf"]).exists():
                missing_pdf = True
            if existing_record.get("cv_docx") and missing_pdf:
                try:
                    existing_record["cv_pdf"] = convert_docx_to_pdf(existing_record["cv_docx"])
                    existing_record["pdf_generation_error"] = None
                    checkpoint_changed = True
                except Exception as exc:
                    existing_record["cv_pdf"] = ""
                    existing_record["pdf_generation_error"] = str(exc)
                    checkpoint_changed = True

            continue

        print(f"Generating docs for job {index}/{total_jobs}: {job_id} - {job.get('title', '')}")

        try:
            generated_payload = generate_docs_for_job(
                deepseek_api_key=deepseek_api_key,
                deepseek_model=args.model,
                gemini_client=gemini_client,
                gemini_model=args.fallback_model,
                cv_text=cv_text,
                job=job,
                candidate_name=candidate_name,
                extra_instructions=args.stage4_extra_prompt,
                prompt_override=args.stage4_prompt_override,
                retries=max(1, args.retries),
                retry_sleep=max(0.0, args.retry_sleep),
            )

            temp_record = {
                **job,
                **generated_payload,
                "run_date": run_date,
                "run_timestamp": run_timestamp,
            }
            ensure_structured_cv_fields(temp_record, candidate_name=candidate_name, cv_text=cv_text)
            cv_doc_path = create_cv_document(
                temp_record,
                docs_dir=docs_dir,
                run_date=run_date,
                candidate_name=candidate_name,
                candidate_email=candidate_email,
                cv_font_name=cv_font_name,
                languages=languages,
                profile_image_path=profile_image_path,
            )
            try:
                cv_pdf_path = convert_docx_to_pdf(cv_doc_path)
                pdf_generation_error = None
            except Exception as exc:
                cv_pdf_path = ""
                pdf_generation_error = str(exc)

            generated_record = {
                **temp_record,
                "cv_docx": cv_doc_path,
                "cv_pdf": cv_pdf_path,
                "tailored_cv_docx": cv_doc_path,
                "pdf_generation_error": pdf_generation_error,
                "doc_generation_error": None,
            }
        except Exception as exc:
            generated_record = {
                **job,
                "cv_professional_summary": "",
                "cv_professional_experience": [],
                "cv_strategic_initiatives": [],
                "cv_skills": [],
                "cv_education": [],
                "tailored_cv": "",
                "cv_docx": "",
                "cv_pdf": "",
                "tailored_cv_docx": "",
                "pdf_generation_error": "",
                "run_date": run_date,
                "run_timestamp": run_timestamp,
                "doc_generation_error": str(exc),
            }

        if not generated_record.get("linkedin_link"):
            generated_record["linkedin_link"] = generated_record.get("link", "")
        if not generated_record.get("apply_link"):
            generated_record["apply_link"] = generated_record.get("linkedin_link", "")

        generated_by_id[job_id] = generated_record
        save_json_file(checkpoint_path, {"generated_records": list(generated_by_id.values())})
        checkpoint_changed = False

        if args.sleep_seconds > 0:
            time.sleep(args.sleep_seconds)

    if checkpoint_changed:
        save_json_file(checkpoint_path, {"generated_records": list(generated_by_id.values())})

    final_records = []
    for job in jobs:
        job_id = str(job.get("job_id"))
        if job_id in generated_by_id:
            final_records.append(generated_by_id[job_id])

    output_json_path = Path(args.output_json)
    output_xlsx_path = Path(args.output_xlsx)
    save_json_file(output_json_path, final_records)
    save_to_excel(
        records=final_records,
        output_path=output_xlsx_path,
        excel_mode=args.excel_mode,
        sheet_name=args.sheet_name.strip(),
        run_date=run_date,
    )

    failed_count = sum(1 for item in final_records if item.get("doc_generation_error"))
    pdf_failed_count = sum(1 for item in final_records if item.get("pdf_generation_error"))
    print("Stage 4 complete.")
    print(f"Generated records: {len(final_records)} -> {output_json_path}")
    print(f"Excel export: {output_xlsx_path} (mode={args.excel_mode})")
    print(f"Word docs directory: {docs_dir.resolve()}")
    print(f"Candidate name: {candidate_name}")
    print(f"Profile image: {profile_image_path if profile_image_path else 'not provided'}")
    print(f"Generation errors: {failed_count}")
    print(f"PDF conversion errors: {pdf_failed_count}")
    print(f"Checkpoint saved: {checkpoint_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
